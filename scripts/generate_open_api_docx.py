from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "开放API接口简化设计.md"
TARGET_DOCX = ROOT / "docs" / "开放API接口文档.docx"


def set_font(run, name: str, size: int, *, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def style_paragraph(paragraph, *, before=0, after=0, line=1.15, left_indent=0, first_line=0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if left_indent:
        fmt.left_indent = Inches(left_indent)
    if first_line:
        fmt.first_line_indent = Inches(first_line)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, after=6)
    run = p.add_run("开放 API 接口文档")
    set_font(run, "Calibri", 24, bold=True, color="1F4E79")

    p2 = doc.add_paragraph()
    style_paragraph(p2, after=18)
    run2 = p2.add_run("治理分析开放接口（3/4/5）简化设计与使用说明")
    set_font(run2, "Calibri", 13, color="44546A")

    meta = doc.add_paragraph()
    style_paragraph(meta, after=18)
    run3 = meta.add_run("版本：2026-07-15    适用范围：/api/v1/public/governance/*")
    set_font(run3, "Calibri", 10, color="666666")


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    set_cell_width(table.cell(0, 0), 6.2)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "EEF4FB")
    p = cell.paragraphs[0]
    style_paragraph(p, before=4, after=4)
    run = p.add_run(text)
    set_font(run, "Calibri", 10, color="1F1F1F")
    doc.add_paragraph()


def add_heading(doc: Document, level: int, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(
        p,
        before=16 if level == 1 else 12 if level == 2 else 8,
        after=6 if level < 3 else 4,
    )
    run = p.add_run(text)
    color = "2E74B5" if level in (1, 2) else "1F4D78"
    size = 16 if level == 1 else 13 if level == 2 else 12
    set_font(run, "Calibri", size, bold=True, color=color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4)
    run = p.add_run(text)
    set_font(run, "Calibri", 10)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=4)
    run = p.add_run(text)
    set_font(run, "Calibri", 10)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, after=6)
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_font(run, "Consolas", 9)
        else:
            run = p.add_run(part)
            set_font(run, "Calibri", 10)


def add_code_block(doc: Document, lines: list[str], language: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    set_cell_width(table.cell(0, 0), 6.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    p = cell.paragraphs[0]
    style_paragraph(p, before=4, after=4)
    if language:
        tag = p.add_run(language.upper() + "\n")
        set_font(tag, "Consolas", 8, color="666666")
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_font(run, "Consolas", 9, color="1F1F1F")
        if idx < len(lines) - 1:
            run.add_break()
    doc.add_paragraph()


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for raw in lines:
        cells = [cell.strip() for cell in raw.strip().strip("|").split("|")]
        rows.append(cells)
    header = rows[0]
    body = [row for row in rows[2:] if any(item for item in row)]
    return header, body


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    col_count = len(header)
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.allow_autofit = False
    widths = {
        2: [2.1, 4.1],
        3: [1.6, 1.6, 3.1],
        4: [1.2, 1.2, 1.4, 2.7],
        5: [1.25, 0.9, 1.0, 0.95, 2.4],
        6: [1.25, 1.0, 1.15, 1.0, 1.05, 1.05],
    }.get(col_count, [6.2 / col_count] * col_count)
    hdr = table.rows[0].cells
    for idx, text in enumerate(header):
        set_cell_width(hdr[idx], widths[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[idx], "D9E2F3")
        p = hdr[idx].paragraphs[0]
        style_paragraph(p, before=2, after=2)
        r = p.add_run(text)
        set_font(r, "Calibri", 9, bold=True, color="1F1F1F")

    for row in rows:
        cells = table.add_row().cells
        for idx in range(col_count):
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[idx] if idx < len(row) else ""
            p = cells[idx].paragraphs[0]
            style_paragraph(p, before=2, after=2)
            r = p.add_run(text)
            set_font(r, "Calibri", 9)
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    add_title(doc)
    add_callout(
        doc,
        "本文档同步自《开放 API 接口简化设计》，用于对外发布治理分析开放接口（3/4/5）的简化契约、默认策略与示例返回。",
    )

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_buffer: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines, code_lang)
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lang = stripped[3:].strip()
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_buffer = [line]
            i += 1
            while i < len(lines):
                nxt = lines[i].rstrip()
                if nxt.strip().startswith("|") and "|" in nxt.strip()[1:]:
                    table_buffer.append(nxt)
                    i += 1
                else:
                    break
            header, body = parse_table(table_buffer)
            add_table(doc, header, body)
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, 1, stripped[3:].strip())
        elif stripped.startswith("### "):
            add_heading(doc, 2, stripped[4:].strip())
        elif stripped.startswith("#### "):
            add_heading(doc, 3, stripped[5:].strip())
        elif re.match(r"^\d+\.\s+", stripped):
            add_numbered(doc, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:].strip())
        else:
            add_paragraph(doc, stripped)
        i += 1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(footer)
    run = footer.add_run("WindEye 开放接口文档")
    set_font(run, "Calibri", 9, color="666666")
    run.add_break(WD_BREAK.LINE)
    run2 = footer.add_run("治理分析开放接口（3/4/5）")
    set_font(run2, "Calibri", 8, color="888888")

    return doc


def main() -> None:
    doc = build_doc()
    TARGET_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET_DOCX)
    print(f"Generated: {TARGET_DOCX}")


if __name__ == "__main__":
    main()
