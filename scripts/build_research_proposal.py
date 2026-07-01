from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "2026CX_陈森_多智能体合规治理研究_修订版.docx"


TITLE = "基于事理图谱的多智能体合规治理研究"
SUBTITLE = "江苏省研究生科研实践创新计划项目申请书"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for p in candidates:
        if p.exists():
            return ImageFont.truetype(str(p), size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline="#6B7280", radius=16, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered(draw, xy, text, fnt, fill="#111827", line_gap=6):
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + line_gap * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + line_gap


def arrow(draw, start, end, color="#334155", width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - sign * 14, y2 - 8), (x2 - sign * 14, y2 + 8)]
    else:
        sign = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 8, y2 - sign * 14), (x2 + 8, y2 - sign * 14)]
    draw.polygon(pts, fill=color)


def make_framework_fig(path: Path):
    img = Image.new("RGB", (1600, 930), "white")
    d = ImageDraw.Draw(img)
    title_f = font(38, True)
    h_f = font(27, True)
    b_f = font(22)
    d.text((60, 36), "图1 研究总体框架：法规-事理图谱-多智能体-协同治理闭环", font=title_f, fill="#0B2545")
    boxes = [
        ((70, 150, 360, 330), "#E0F2FE", "资本市场 AI 应用场景\n投顾/交易/风控/披露\n客服/研报/合规审查"),
        ((70, 390, 360, 570), "#ECFDF5", "多源治理数据\n法规条款、监管案例\n公告舆情、企业事件"),
        ((470, 235, 760, 485), "#F8FAFC", "事理知识图谱\n主体-事件-特征-法规\n责任-证据-处置措施"),
        ((870, 130, 1180, 300), "#FEF3C7", "多智能体协同\n检索/风险/合规\n证据校验/共识聚合"),
        ((870, 365, 1180, 555), "#FCE7F3", "证据链约束\n图谱路径 + 法规依据\n来源记录 + 置信度"),
        ((1270, 235, 1530, 485), "#EDE9FE", "治理输出\n分级监管策略\n穿透式协同方案\n报告与反馈更新"),
    ]
    for xy, fill, text in boxes:
        rounded_box(d, xy, fill)
        centered(d, xy, text, h_f if "事理" in text else b_f)
    arrow(d, (360, 240), (470, 300))
    arrow(d, (360, 480), (470, 420))
    arrow(d, (760, 320), (870, 215))
    arrow(d, (760, 400), (870, 455))
    arrow(d, (1180, 215), (1270, 310))
    arrow(d, (1180, 455), (1270, 395))
    d.arc((440, 650, 1370, 890), 0, 180, fill="#2563EB", width=5)
    arrow(d, (470, 770), (360, 560), "#2563EB")
    centered(d, (610, 675, 1200, 755), "人工复核与处置结果回流：修正图谱、阈值、规则与智能体提示", b_f, fill="#1D4ED8")
    img.save(path)


def make_agent_fig(path: Path):
    img = Image.new("RGB", (1600, 930), "white")
    d = ImageDraw.Draw(img)
    title_f = font(38, True)
    b_f = font(22)
    small_f = font(19)
    d.text((60, 36), "图2 M1-M4 多智能体协同治理架构", font=title_f, fill="#0B2545")
    xs = [80, 450, 820, 1190]
    labels = [
        "M1 感知与检索层\n意图识别 / 实体对齐\nRAG 召回 / 5级图查询\n多源证据并行提取",
        "M2 图谱推理层\n社区发现 / 路径枚举\n风险传导评分\n异常结构识别",
        "M3 合规校验层\n法规条款匹配\n三级指标评分\n证据链完整性校验",
        "M4 共识治理层\n冲突消解 / 置信聚合\n分级监管建议\n报告生成与反馈回流",
    ]
    colors = ["#DBEAFE", "#FEE2E2", "#DCFCE7", "#EDE9FE"]
    for x, label, color in zip(xs, labels, colors):
        xy = (x, 170, x + 290, 405)
        rounded_box(d, xy, color)
        centered(d, xy, label, b_f)
    for x in xs[:-1]:
        arrow(d, (x + 290, 285), (x + 370, 285))
    rounded_box(d, (320, 520, 1280, 705), "#F8FAFC", "#94A3B8")
    centered(
        d,
        (350, 535, 1250, 690),
        "统一上下文 GovernanceContext：query、resolved_entities、subgraph、risk_paths、compliance、scoring、governance\n"
        "统一证据链 EvidenceChain：结论、图谱路径、法规条款、数据来源、时间戳、置信度、责任主体",
        small_f,
    )
    arrow(d, (1335, 405), (1200, 520), "#7C3AED")
    arrow(d, (470, 520), (280, 405), "#7C3AED")
    rounded_box(d, (480, 765, 1120, 855), "#ECFDF5", "#059669")
    centered(d, (500, 775, 1100, 845), "输出：可信合规结论、分级监管建议、协同处置清单、可追溯治理报告", b_f, fill="#065F46")
    img.save(path)


def make_route_fig(path: Path):
    img = Image.new("RGB", (1600, 930), "white")
    d = ImageDraw.Draw(img)
    title_f = font(38, True)
    b_f = font(22)
    d.text((60, 36), "图3 技术路线：从场景建模到原型验证", font=title_f, fill="#0B2545")
    lanes = [
        ("阶段一：真实数据接入", "深交所/北交所纪律处分\n监管处罚与问询函\n上市公司公告\n法规条款与案例文本"),
        ("阶段二：四层图谱构建", "主体/事件/特征/法规抽取\n法规条款链接\n证据链结构化\nNeo4j 图谱入库"),
        ("阶段三：M1-M4协同推理", "RAG并行提取\n5级Cypher穿透查询\n路径评分与合规校验\n共识治理报告"),
        ("阶段四：实验评价", "真实案例库与基线模型\n消融实验\n专家评分\n原型展示与报告导出"),
    ]
    x = 80
    for idx, (head, body) in enumerate(lanes):
        xy = (x, 190, x + 320, 650)
        fill = ["#EFF6FF", "#F0FDFA", "#FFF7ED", "#F5F3FF"][idx]
        rounded_box(d, xy, fill)
        centered(d, (x + 20, 210, x + 300, 285), head, b_f, fill="#0F172A")
        centered(d, (x + 20, 320, x + 300, 610), body, b_f)
        if idx < len(lanes) - 1:
            arrow(d, (x + 320, 420), (x + 390, 420))
        x += 380
    d.line((120, 760, 1480, 760), fill="#64748B", width=5)
    for i, txt in enumerate(["M1-M3", "M4-M6", "M7-M10", "M11-M12"]):
        cx = 210 + i * 380
        d.ellipse((cx - 16, 744, cx + 16, 776), fill="#2563EB")
        centered(d, (cx - 80, 790, cx + 80, 835), txt, b_f, fill="#1E3A8A")
    img.save(path)


def make_eval_fig(path: Path):
    img = Image.new("RGB", (1600, 930), "white")
    d = ImageDraw.Draw(img)
    title_f = font(38, True)
    b_f = font(22)
    small_f = font(19)
    d.text((60, 36), "图4 试验方案：四类基线、五组指标、三类案例", font=title_f, fill="#0B2545")
    sections = [
        ((80, 155, 470, 455), "#F8FAFC", "对比基线", "B1 单一 LLM 直接研判\nB2 文档 RAG\nB3 图谱 RAG\nB4 无共识多智能体"),
        ((605, 155, 995, 455), "#F8FAFC", "评价指标", "准确率/F1\n证据覆盖率\n幻觉率\n治理建议可执行性\n响应时延与成本"),
        ((1130, 155, 1520, 455), "#F8FAFC", "案例场景", "AI 投顾适当性\n算法交易异常\nAI 研报/披露合规"),
    ]
    for xy, fill, head, body in sections:
        rounded_box(d, xy, fill)
        centered(d, (xy[0] + 20, xy[1] + 25, xy[2] - 20, xy[1] + 90), head, b_f, fill="#1F4D78")
        centered(d, (xy[0] + 30, xy[1] + 115, xy[2] - 30, xy[3] - 20), body, small_f)
    arrow(d, (470, 305), (605, 305))
    arrow(d, (995, 305), (1130, 305))
    rounded_box(d, (225, 585, 1375, 780), "#ECFDF5", "#059669")
    centered(
        d,
        (245, 600, 1355, 760),
        "核心验证假设：引入“事理图谱 + 证据链校验 + 共识聚合”后，\n"
        "合规结论准确性和可追溯性提升，幻觉率下降，治理建议更具可执行性。",
        b_f,
        fill="#064E3B",
    )
    img.save(path)


def set_cell(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        rest = text[len(bold_prefix):]
        if rest:
            rr = p.add_run(rest)
            rr.font.name = "宋体"
            rr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    else:
        r = p.add_run(text)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    r.italic = True


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, size=9)
        shade_cell(hdr[i], "F4F6F9")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], str(val), size=9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def build_doc():
    fig1 = OUT_DIR / "fig1_framework.png"
    fig2 = OUT_DIR / "fig2_agents.png"
    fig3 = OUT_DIR / "fig3_route.png"
    fig4 = OUT_DIR / "fig4_eval.png"
    make_framework_fig(fig1)
    make_agent_fig(fig2)
    make_route_fig(fig3)
    make_eval_fig(fig4)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        st = styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = RGBColor(46, 116, 181)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(20)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目类别建议：人文社科项目（交叉融合信息科学方法）")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(11)
    r.italic = True

    add_table(
        doc,
        ["字段", "建议填写"],
        [
            ["项目名称", TITLE],
            ["起止年限", "2026年7月至2027年6月（可按学校通知调整）"],
            ["研究方向", "资本市场人工智能应用监管的多元主体协同治理方法"],
            ["依托基础", "WindEye 事理知识图谱与 DRA-MA 多智能体协同治理原型"],
            ["拟形成成果", "论文1篇、原型系统1套、案例库/知识图谱1份、研究报告1份、申请并获批软件著作权1项"],
        ],
        widths=[3.0, 13.0],
    )

    doc.add_heading("一、立项依据", level=1)
    add_para(
        doc,
        "本项目拟围绕资本市场人工智能应用监管中的权责边界不清、风险证据分散、治理主体协同不足和大模型生成结论可信度不足等问题，"
        "开展“基于事理图谱的多智能体合规治理研究”。项目来源于申请人已有的 WindEye 平台实践基础：该平台已形成数据采集、四层事理知识图谱、"
        "风险路径分析、合规评分、治理报告生成等能力，为本项目开展可验证的研创工作提供了工程和数据基础。"
    )
    add_para(
        doc,
        "从政策背景看，生成式人工智能服务在金融投顾、算法交易、智能客服、研报生成、合规审查、风控预警等场景中快速渗透，"
        "一方面提升资本市场信息处理效率，另一方面也带来虚假信息、算法偏见、数据合规、模型幻觉、责任归属和系统性风险放大等问题。"
        "我国《生成式人工智能服务管理暂行办法》强调促进发展与规范应用并重，并要求提高生成内容准确性、可靠性和透明度；"
        "《人工智能安全治理框架》1.0版提出从模型算法、数据、系统和应用等维度开展风险管理；欧盟 AI Act 采用风险分级监管逻辑，"
        "要求高风险 AI 系统具备数据治理、日志记录、透明度、人类监督和准确性等机制。这些制度趋势表明，资本市场 AI 应用监管正在从"
        "事后合规审查转向分级动态监管、穿透式证据审查和多主体协同治理。"
    )
    add_para(
        doc,
        "从现实需求看，资本市场监管对象具有强关联、强时序和强传导特征。AI 投顾可能涉及投资者适当性和误导性推荐；算法交易可能引发异常波动、"
        "市场操纵或风险联动；AI 研报和披露辅助工具可能产生事实性错误或法规引用错误。传统规则库和单一大模型均难以同时满足“看清关系、说明证据、"
        "匹配规则、生成治理方案”的要求。因此，需要将事理知识图谱的结构化证据能力、多智能体的任务分工能力和合规规则的约束能力结合起来，"
        "形成面向资本市场 AI 应用监管的可解释、可追溯、可复核的协同治理方法。"
    )
    doc.add_picture(str(fig1), width=Inches(6.5))
    add_caption(doc, "图1 研究总体框架")

    doc.add_heading("1.1 国内外研究现状", level=2)
    add_para(doc, "（1）人工智能监管与金融 RegTech 研究。")
    add_para(
        doc,
        "国际研究普遍认为，AI 在金融服务中的应用正在改变监管对象、风险结构和治理方式。美国 GAO 2025 年报告指出，金融监管机构已经开始使用 AI 工具"
        "开展市场监督和被监管主体监督，同时 AI 应用也带来模型风险、数据风险和监督能力建设问题。IMF 2025 年针对证券市场 AI 加速应用的技术报告"
        "进一步指出，生成式 AI 在证券市场中的应用会带来数据、性能、网络安全和金融稳定风险。学术界关于 RegTech、XAI in Finance、AI Governance 的研究"
        "强调自动化合规、可解释模型、审计追踪和监管沙箱，但现有研究仍较少把“法规条款-风险事件-市场主体-治理动作”作为统一事理链条进行建模。"
    )
    add_para(doc, "（2）知识图谱与事理图谱研究。")
    add_para(
        doc,
        "知识图谱能够以实体、关系、事件和属性的形式表示复杂对象，是金融风险穿透识别和合规证据组织的重要基础。近年来，大模型与知识图谱融合成为热点："
        "Think-on-Graph 将大模型视为图谱上的交互式推理智能体，通过图搜索缓解深度推理中的幻觉；ACL、EMNLP 2024 多项研究围绕路径选择、KG-Adapter、"
        "KG 规划数据和 KGQA 检索推理展开，说明结构化图谱对复杂问答和事实推理具有明显价值。与一般知识图谱相比，事理知识图谱更强调事件因果、风险传导、"
        "时间演化和处置后果，适合表达资本市场 AI 应用监管中的“行为-风险-责任-措施”链条。"
    )
    add_para(doc, "（3）多智能体协同与大模型可信推理研究。")
    add_para(
        doc,
        "多智能体大模型系统已从简单角色扮演走向流程化协作和任务分解。CAMEL、MetaGPT、ChatDev 等工作表明，多个具备角色分工的智能体能够通过通信和"
        "流程约束完成复杂任务；其中 MetaGPT 特别指出，朴素串联大模型会产生级联幻觉，需要将标准作业流程和中间结果校验嵌入协作过程。Self-RAG、"
        "Corrective RAG 和 GraphRAG 等研究进一步说明，检索、反思、图结构索引和社区摘要有助于提升生成事实性和可追溯性。"
    )
    add_para(doc, "（4）不足与发展趋势。")
    add_bullets(
        doc,
        [
            "现有 AI 治理研究多停留在制度框架和风险分类层面，缺少可计算的证据链、责任链和治理动作链建模。",
            "现有金融合规技术多依赖规则库或单模型问答，对跨主体、跨事件、跨法规的风险传导和监管协同支持不足。",
            "现有多智能体系统虽能分工协作，但在高风险合规场景中仍缺少强证据约束、置信度聚合和人工复核闭环。",
            "未来趋势是分级动态监管、监管科技与合规科技融合、知识图谱增强大模型、可审计 AI 以及人机协同治理闭环。",
        ],
    )

    doc.add_heading("二、研究目标、研究内容和拟解决的主要问题", level=1)
    doc.add_heading("2.1 研究目标", level=2)
    add_para(
        doc,
        "本项目面向资本市场人工智能应用监管场景，构建基于事理知识图谱的多智能体协同合规治理方法。总体目标是："
        "以法规条款、市场主体、AI 应用行为、风险事件、证据来源和治理措施为核心对象，构建可追溯事理图谱；"
        "设计由图谱检索、风险分析、合规研判、证据校验和共识聚合等智能体组成的协同推理机制；"
        "形成分级动态监管策略和多元主体穿透式协同治理策略；并依托 WindEye 平台完成原型验证与案例评估。"
    )
    add_numbered(
        doc,
        [
            "构建资本市场 AI 应用合规治理事理图谱本体，表达“应用场景-主体角色-风险事件-法规条款-证据链-治理措施”。",
            "设计证据链约束的多智能体协同推理机制，使每一项合规结论均可回溯到图谱路径、法规依据和数据来源。",
            "提出面向资本市场 AI 应用的分级动态监管策略，依据风险等级、证据充分性和主体责任差异生成治理建议。",
            "开发可运行原型并完成案例验证，比较单一 LLM、普通 RAG、图谱 RAG 和本项目方法在准确性、可追溯性和幻觉率上的差异。",
        ],
    )

    doc.add_heading("2.2 研究内容", level=2)
    add_table(
        doc,
        ["研究内容", "核心任务", "预期产出"],
        [
            ["内容一：合规治理事理图谱构建", "设计资本市场 AI 应用监管本体；抽取主体、事件、风险特征、法规条款和治理措施；建立证据链对象。", "本体模型、样例图谱、数据字典"],
            ["内容二：M1-M4 多智能体协同推理", "构建M1感知检索、M2图谱推理、M3合规校验、M4共识治理四个模块，并设计统一上下文、消息流和状态流。", "多智能体推理框架、共识规则"],
            ["内容三：分级动态监管与穿透式协同治理", "根据风险路径、主体角色、社区结构和法规匹配结果生成分级监管策略、责任分配和处置清单。", "治理策略模型、处置模板"],
            ["内容四：原型系统与实证评价", "基于 WindEye 平台实现原型，构建案例集，开展基线对比、消融实验和专家评价。", "原型系统、实验报告、论文"],
        ],
        widths=[4.0, 8.0, 4.0],
    )
    add_para(
        doc,
        "其中，多智能体部分不采用松散的角色对话形式，而采用 M1-M4 的工程化协同架构。M1 感知与检索模块负责意图识别、实体对齐、RAG 证据召回和多源文本并行提取；"
        "M2 图谱推理模块负责在 Neo4j 中进行多跳子图检索、风险社区发现和传导路径枚举；M3 合规校验模块负责法规条款匹配、三级合规指标评分和证据链完整性检查；"
        "M4 共识治理模块负责冲突消解、置信度聚合、分级监管建议生成和治理报告输出。四个模块通过统一的 GovernanceContext 传递实体、子图、路径、合规结果、"
        "评分和治理建议，避免各环节各自为政。"
    )

    doc.add_heading("2.3 拟解决的主要问题", level=2)
    add_bullets(
        doc,
        [
            "监管对象难以穿透建模：将 AI 应用、数据来源、算法行为、服务对象、市场主体和监管规则纳入统一事理图谱。",
            "合规结论缺乏证据约束：设计 EvidenceChain 对象，使每个结论具备路径、条款、来源、时间和置信度。",
            "多智能体协同易产生冲突和级联幻觉：通过证据覆盖率、一致性投票、低置信拦截和人工复核机制提高可靠性。",
            "治理建议难以落实到主体责任：根据监管部门、交易所、市场机构、AI 服务提供者和投资者等角色生成协同处置方案。",
            "研究成果难以验证：构建案例库和指标体系，用可量化实验评估准确率、幻觉率、证据覆盖率和治理建议可执行性。",
        ],
    )

    doc.add_heading("2.4 与本人学位论文/实践成果之间的关系", level=2)
    add_para(
        doc,
        "本项目与申请人的学位论文和实践成果具有直接延续关系。申请人已参与 WindEye 平台建设，围绕事理知识图谱构建、知识图谱查询分析、风险路径识别、"
        "合规评分和协同治理报告生成开展工程实践。该实践成果已经具备数据采集、Neo4j 图谱存储、AntV G6 可视化、DRA-MA 多阶段分析管线、SSE 流式推理展示等基础。"
        "本研创项目将在已有系统基础上进一步深化理论模型和实验评价，重点把实践中的工程模块上升为“事理图谱驱动的多智能体合规治理方法”，"
        "为学位论文中的知识图谱推理、智能治理和金融监管科技研究提供核心章节、实验平台和案例材料。"
    )

    doc.add_heading("三、研究思路与方法、技术路线、试验方案及可行性分析", level=1)
    doc.add_heading("3.1 研究思路", level=2)
    add_para(
        doc,
        "本项目遵循“场景牵引、图谱表达、智能体协同、证据约束、治理闭环”的研究思路。首先，从资本市场 AI 应用监管场景出发，梳理 AI 投顾、算法交易、"
        "智能研报、合规审查等典型业务中的风险事件和监管要求；其次，构建事理知识图谱，将市场主体、AI 系统、数据来源、算法行为、风险特征、法规条款、"
        "证据来源和治理措施统一表达；再次，设计 M1-M4 多智能体协同机制，由感知检索、图谱推理、合规校验和共识治理四个模块依次完成证据召回、路径解释、规则匹配和治理输出；"
        "最后，通过专家反馈和案例验证形成动态更新的治理闭环。"
    )
    doc.add_picture(str(fig2), width=Inches(6.5))
    add_caption(doc, "图2 M1-M4 多智能体协同治理架构")

    doc.add_heading("3.2 研究方法", level=2)
    add_para(doc, "（1）事理知识图谱建模方法。")
    add_para(
        doc,
        "采用“主体-事件-风险特征-法规-治理动作”的多层本体结构。主体层包括监管机构、交易所、证券基金经营机构、上市公司、AI 服务提供者、投资者等；"
        "事件层包括 AI 推荐、算法交易、信息披露、模型更新、异常波动、投诉举报、监管处罚等；风险特征层包括数据来源不明、推荐误导、适当性不匹配、"
        "模型幻觉、算法操纵、信息披露错误等；法规层包括法律法规、监管规则、行业标准和内部制度；治理动作层包括风险提示、人工复核、模型下线、"
        "补充披露、升级监管、持续监测等。"
    )
    add_para(
        doc,
        "图谱穿透查询采用 Neo4j 5 级深度路径逻辑实现：第1级定位种子主体或 AI 应用事件；第2级扩展直接关联主体，包括投资、担保、任职、控制、共同涉案等关系；"
        "第3级连接风险事件和异常行为，如纪律处分、监管问询、负面舆情、异常交易和信息披露问题；第4级映射风险特征，包括数据来源、算法偏差、内容真实性、"
        "适当性匹配、市场操纵风险等；第5级链接法规条款和治理动作，形成“主体—关系—事件—特征—法规/处置”的完整证据链。"
        "在实现上，可通过 Cypher 的 MATCH p=(seed)-[*1..5]-(n) 进行候选路径扩展，并结合关系白名单、节点标签过滤、路径长度惩罚和风险权重排序，"
        "从表层关联中筛选出隐蔽但具有合规意义的风险链条。"
    )
    add_para(doc, "（2）多智能体协同推理方法。")
    add_para(
        doc,
        "采用 M1-M4 分层架构组织多智能体协同推理。M1 感知与检索层面向用户问题、上传材料和监管文本，完成意图识别、命名实体识别、实体消歧、"
        "RAG 召回和并行证据提取；M2 图谱推理层基于四层事理图谱执行 k-hop 子图扩展、社区发现、中心性计算、风险路径枚举和异常结构识别；"
        "M3 合规校验层将风险路径与法规条款、合规指标和证据来源进行交叉验证，输出合规匹配、节点合规分和证据链完整度；"
        "M4 共识治理层对来自 M2、M3 的结论进行一致性判断、置信度聚合和冲突消解，生成分级监管建议、协同处置清单和最终治理报告。"
        "四个模块通过结构化消息传递，统一写入 GovernanceContext，保证每一步结果均可审计、可复核。"
    )
    add_para(doc, "（3）证据链约束与共识聚合方法。")
    add_para(
        doc,
        "将每个候选结论表示为 EvidenceChain = <claim, graph_path, regulation_clause, source, timestamp, confidence, responsible_subject>。"
        "其中 claim 为风险或合规结论，graph_path 为支撑该结论的图谱路径，regulation_clause 为匹配法规条款，source 为数据来源，confidence 为综合置信度。"
        "共识分数可由证据覆盖率、图谱路径有效性、法规匹配置信度、智能体一致性和数据源可信度加权得到："
    )
    add_para(
        doc,
        "ConsensusScore = w1*PathValidity + w2*RegulationMatch + w3*SourceReliability + w4*AgentAgreement + w5*EvidenceCompleteness。"
    )
    add_para(
        doc,
        "当 ConsensusScore 高于阈值时输出可信合规结论；处于中间区间时输出“需人工复核”；低于阈值时标记为证据不足或结论冲突，从而降低大模型幻觉对治理结论的影响。"
    )
    doc.add_picture(str(fig3), width=Inches(6.5))
    add_caption(doc, "图3 技术路线")

    doc.add_heading("3.3 技术路线", level=2)
    add_numbered(
        doc,
        [
            "需求与本体设计：梳理资本市场 AI 应用场景、监管规则和典型风险事件，形成合规治理本体和数据字典。",
            "数据采集与图谱构建：采集深交所、北交所、上交所纪律处分/监管问询/自律监管措施、证监会处罚决定、上市公司公告和法规政策文本；进行实体抽取、事件抽取、关系抽取和法规链接；导入 Neo4j。",
            "智能体编排与推理：基于现有 DRA-MA 管线，将 M1 感知检索、M2 图谱推理、M3 合规校验、M4 共识治理封装为可替换模块，并通过统一上下文串联。",
            "五级穿透查询：在 Neo4j 中执行 1-5 跳 Cypher 查询，按主体、关系、事件、风险特征、法规/处置五类节点逐级扩展，并结合关系白名单与风险权重过滤隐蔽风险链。",
            "证据链与共识机制：实现 EvidenceChain 结构，设计证据覆盖率、路径有效性、法规匹配和多智能体一致性评分。",
            "分级监管策略生成：依据风险等级、证据充分性、主体责任和处置紧急度生成分级动态监管策略和穿透式协同治理方案。",
            "试验验证与优化：构建案例集，与单一 LLM、普通 RAG、图谱 RAG、无证据共识多智能体等基线比较，开展消融实验和专家评分。",
        ],
    )

    doc.add_heading("3.4 试验方案", level=2)
    add_para(
        doc,
        "试验设计坚持“可复现、可比较、可解释”的原则，围绕三类资本市场 AI 应用监管场景开展：AI 投顾适当性风险、算法交易异常风险、AI 研报/披露合规风险。"
        "数据来源以真实公开数据为主，包括深圳证券交易所纪律处分与监管措施、北京证券交易所纪律处分和自律监管措施、上海证券交易所监管问询与纪律处分、"
        "中国证监会行政处罚和市场禁入决定、上市公司公告、公开新闻舆情文本以及项目构造的模拟 AI 应用场景。"
    )
    add_table(
        doc,
        ["试验维度", "设计说明"],
        [
            ["数据集构建", "不少于100个案例片段，覆盖法规条款、主体关系、风险事件、处置动作；人工标注实体、关系、风险等级、法规依据和标准结论。"],
            ["对比基线", "单一 LLM 直接生成、文档 RAG、图谱 RAG、无证据校验多智能体、本项目完整方法。"],
            ["评价指标", "实体/关系抽取 F1、法规匹配准确率、风险结论准确率、证据覆盖率、幻觉率、治理建议可执行性、响应时延和人工复核率。"],
            ["消融实验", "去除事理图谱、去除合规智能体、去除证据校验、去除共识聚合，比较各模块贡献。"],
            ["专家评价", "邀请相关专业教师或金融/合规从业人员，从准确性、可解释性、可操作性、监管适配性四方面评分。"],
        ],
        widths=[3.2, 12.8],
    )
    doc.add_picture(str(fig4), width=Inches(6.5))
    add_caption(doc, "图4 试验方案")

    doc.add_heading("3.5 创新性", level=2)
    add_bullets(
        doc,
        [
            "提出资本市场 AI 应用监管的事理图谱表达框架，将 AI 应用场景、市场主体、风险事件、法规条款和治理措施纳入统一语义结构。",
            "提出证据链约束的多智能体合规共识推理机制，避免单一大模型直接输出高风险合规结论，提高可追溯性和可信度。",
            "提出面向资本市场 AI 应用的分级动态监管策略，根据风险等级、证据充分性、主体责任和协同关系生成差异化治理方案。",
            "构建从图谱检索到治理报告的闭环实验系统，使制度研究、算法方法和工程原型相互支撑，增强研创项目落地性。",
            "设计多维评价体系，将幻觉率、证据覆盖率、路径有效性和治理建议可执行性纳入实验指标，补足传统准确率评价不足。",
        ],
    )

    doc.add_heading("3.6 可行性分析", level=2)
    add_para(
        doc,
        "（1）技术基础可行。WindEye 已具备 FastAPI、Neo4j、React、AntV G6、知识图谱查询、社区发现、风险路径分析、合规评分和报告导出等模块，"
        "可直接作为研创项目原型底座。"
    )
    add_para(
        doc,
        "（2）数据来源可行。项目使用公开、可核验的一手监管数据和市场文本，重点包括深圳证券交易所纪律处分报告、北京证券交易所纪律处分与自律监管措施、"
        "上海证券交易所监管问询和纪律处分、中国证监会行政处罚决定、上市公司公告、公开舆情事件和政策法规文本，避免依赖难以获取的敏感内部数据；"
        "对智能投顾、AI 研报生成等新兴应用场景，可在真实监管逻辑基础上构造少量模拟案例补充。"
    )
    add_para(
        doc,
        "（3）研究周期可行。项目可在一年内完成本体设计、图谱构建、智能体封装、共识机制实现、案例验证和论文撰写。核心功能基于已有系统扩展，研发风险较低。"
    )
    add_para(
        doc,
        "（4）成果形态可行。项目成果既可形成论文和研究报告，也可沉淀为原型系统、知识图谱样例库、治理评估指标和可复用软件组件，符合研究生科研实践创新计划要求。"
    )

    doc.add_heading("四、研究工作的总体安排及进度", level=1)
    add_table(
        doc,
        ["阶段", "时间", "主要任务", "阶段成果"],
        [
            ["第一阶段", "2026.07-2026.09", "完成文献综述、政策梳理、应用场景界定和本体设计。", "研究综述、本体模型、数据清单"],
            ["第二阶段", "2026.10-2026.12", "采集法规、案例和公告文本；完成实体、事件、法规条款和治理措施抽取；构建样例事理图谱。", "样例图谱、数据字典、标注样本"],
            ["第三阶段", "2027.01-2027.03", "实现多智能体协同推理、证据链对象、共识聚合和分级监管策略生成。", "原型模块、算法说明、阶段测试"],
            ["第四阶段", "2027.04-2027.05", "开展基线对比、消融实验、专家评价和案例分析；优化模型和前端展示。", "实验结果、案例报告"],
            ["第五阶段", "2027.06", "完成论文、研究报告、成果整理和项目结题材料。", "论文投稿、研究报告、原型系统演示"],
        ],
        widths=[2.3, 3.0, 7.8, 3.0],
    )

    doc.add_heading("五、研究工作的预期成果及成果提交形式", level=1)
    add_bullets(
        doc,
        [
            "论文成果：撰写与投稿学术论文1篇，主题为事理图谱、多智能体协同推理或资本市场 AI 合规治理。",
            "系统成果：形成“基于事理图谱的多智能体合规治理”原型系统1套，包含图谱检索、风险路径、合规研判、证据链展示和治理报告模块。",
            "数据成果：形成资本市场 AI 应用监管案例库和样例事理知识图谱，包含实体、事件、法规、证据链和治理措施等结构化数据。",
            "报告成果：提交项目研究报告1份，包含研究背景、方法模型、技术路线、实验结果、案例分析和治理建议。",
            "实践成果：申请并获批软件著作权1项，形成可展示的系统演示视频或截图材料。",
        ],
    )

    doc.add_heading("六、研究基础和工作条件", level=1)
    add_para(
        doc,
        "申请人已具备本项目所需的工程实践基础和研究条件。现有 WindEye 项目以“数据采集-四层知识图谱-多智能体协同推理”为核心架构，"
        "后端采用 FastAPI、Neo4j、NetworkX 等技术，前端采用 React、TypeScript、Ant Design Pro 和 AntV G6，已实现知识图谱查询分析、社区发现、"
        "风险传导路径分析、合规指标评分、治理报告生成和 Word 导出等功能。"
    )
    add_para(
        doc,
        "在数据和工具方面，项目可利用公开法规政策、监管案例、上市公司公告、舆情事件文本和模拟 AI 应用场景构建实验数据；可使用 Neo4j 存储图谱，"
        "使用大模型 API 或本地模型完成抽取、推理和报告生成，使用 Python 进行指标计算和实验分析。学校导师、课题组和相关课程训练可为项目提供方法指导、"
        "论文写作和实验评价支持。"
    )
    add_para(
        doc,
        "项目风险主要包括数据标注成本、法规匹配准确性和大模型输出稳定性。应对措施包括：优先采用公开数据和小规模高质量标注；通过证据链和规则约束降低生成不确定性；"
        "设置人工复核机制；通过消融实验明确各模块贡献；在项目周期内控制系统边界，确保按期完成核心成果。"
    )

    doc.add_heading("七、经费概算", level=1)
    add_para(
        doc,
        "经费使用坚持与研究任务直接相关、厉行节约和可核验原则，主要用于文献资料、公开数据整理、实验算力、论文发表、调研交流和成果材料制作。"
        "具体金额可根据学校最终资助额度等比例调整。"
    )
    add_table(
        doc,
        ["经费科目", "预算金额（元）", "用途说明"],
        [
            ["文献资料费", "800", "购买或下载研究所需图书、标准、法规汇编、数据库检索资料及文献整理工具。"],
            ["数据采集与整理费", "1200", "整理深交所、北交所、上交所纪律处分/监管问询/自律监管措施、证监会处罚决定、上市公司公告等公开数据，进行格式转换、标注和清洗。"],
            ["实验算力与服务器费", "1500", "用于 Neo4j 图数据库运行、向量检索、模型调用、实验脚本运行和系统演示部署。"],
            ["论文版面与投稿费", "1200", "用于论文投稿、版面、查重和英文摘要润色等与成果发表直接相关支出。"],
            ["调研交流费", "800", "用于参加校内外学术交流、专家咨询、案例研讨和相关交通通讯支出。"],
            ["材料打印与成果制作费", "500", "用于申请书、结题材料、研究报告、系统截图、演示材料和软著申报材料制作。"],
            ["合计", "6000", "按江苏省研究生科研实践创新计划常见资助额度预估，最终以学校核定经费为准。"],
        ],
        widths=[3.0, 2.5, 10.5],
    )

    doc.add_heading("参考文献", level=1)
    refs = [
        "Hong S, Zhuge M, Chen J, et al. MetaGPT: Meta Programming for A Multi-Agent Collaborative Framework. ICLR, 2024.（CCF-A）",
        "Li G, Hammoud H, Itani H, et al. CAMEL: Communicative Agents for \"Mind\" Exploration of Large Language Model Society. NeurIPS, 2023.（CCF-A）",
        "Qian C, Cong X, Yang C, et al. ChatDev: Communicative Agents for Software Development. ACL, 2024.（CCF-A）",
        "Wu Q, Bansal G, Zhang J, et al. AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation Framework. arXiv, 2023.",
        "Asai A, Wu Z, Wang Y, et al. Self-RAG: Learning to Retrieve, Generate, and Critique Through Self-Reflection. ICLR, 2024.（CCF-A）",
        "Yao S, Zhao J, Yu D, et al. ReAct: Synergizing Reasoning and Acting in Language Models. ICLR, 2023.（CCF-A）",
        "Sun J, Xu C, Tang L, et al. Think-on-Graph: Deep and Responsible Reasoning of Large Language Model on Knowledge Graph. ICLR, 2024.（CCF-A）",
        "Chen Z, Bai L, Li Z, et al. A New Pipeline for Knowledge Graph Reasoning Enhanced by Large Language Models Without Fine-Tuning. EMNLP, 2024.（CCF-B）",
        "Wang J, Chen M, Hu B, et al. Learning to Plan for Retrieval-Augmented Large Language Models from Knowledge Graphs. Findings of EMNLP, 2024.",
        "Ji Y, Wu K, Li J, et al. Retrieval and Reasoning on KGs: Integrate Knowledge Graphs into Large Language Models for Complex Question Answering. Findings of EMNLP, 2024.",
        "Liu H, Wang S, Zhu Y, et al. Knowledge Graph-Enhanced Large Language Models via Path Selection. Findings of ACL, 2024.",
        "Tian Y, Song H, Wang Z, et al. KG-Adapter: Enabling Knowledge Graph Integration in Large Language Models through Parameter-Efficient Fine-Tuning. Findings of ACL, 2024.",
        "Edge D, Trinh H, Cheng N, et al. From Local to Global: A Graph RAG Approach to Query-Focused Summarization. arXiv, 2024.",
        "Gao Y, Xiong Y, Gao X, et al. Retrieval-Augmented Generation for Large Language Models: A Survey. arXiv, 2024.",
        "Yan S, Gu J, Zhu Y, et al. Corrective Retrieval Augmented Generation. arXiv, 2024.",
        "Guo T, Chen X, Wang Y, et al. Large Language Model based Multi-Agents: A Survey of Progress and Challenges. IJCAI, 2024.（CCF-A）",
        "Wang L, Ma C, Feng X, et al. A survey on large language model based autonomous agents. Frontiers of Computer Science, 2024.（SCI）",
        "Hogan A, Blomqvist E, Cochez M, et al. Knowledge Graphs. ACM Computing Surveys, 2021.（CCF-A/SCI）",
        "Ji S, Pan S, Cambria E, et al. A Survey on Knowledge Graphs: Representation, Acquisition, and Applications. IEEE TNNLS, 2022.（SCI一区）",
        "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS, 2020.（CCF-A）",
        "Wei J, Wang X, Schuurmans D, et al. Chain-of-Thought Prompting Elicits Reasoning in Large Language Models. NeurIPS, 2022.（CCF-A）",
        "Wang X, Wei J, Schuurmans D, et al. Self-Consistency Improves Chain of Thought Reasoning in Language Models. ICLR, 2023.（CCF-A）",
        "Mialon G, Dessì R, Lomeli M, et al. Augmented Language Models: A Survey. Transactions on Machine Learning Research, 2023.",
        "Maynez J, Narayan S, Bohnet B, et al. On Faithfulness and Factuality in Abstractive Summarization. ACL, 2020.（CCF-A）",
        "Ji Z, Lee N, Frieske R, et al. Survey of Hallucination in Natural Language Generation. ACM Computing Surveys, 2023.（CCF-A/SCI）",
        "Mitchell M, Wu S, Zaldivar A, et al. Model Cards for Model Reporting. FAT*, 2019.",
        "Gebru T, Morgenstern J, Vecchione B, et al. Datasheets for Datasets. Communications of the ACM, 2021.",
        "Charoenwong B, Kowaleski Z T, Kwan A, et al. RegTech. Journal of Financial Economics, 2024.（金融学顶级期刊）",
        "Goodell J W, Kumar S, Lim W M, et al. Artificial intelligence and machine learning in finance: Identifying foundations, themes, and research clusters from bibliometric analysis. Journal of Behavioral and Experimental Finance, 2021.",
        "Bussmann N, Giudici P, Marinelli D, et al. Explainable Machine Learning in Credit Risk Management. Computational Economics, 2021.",
        "Arrieta A B, Díaz-Rodríguez N, Del Ser J, et al. Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. Information Fusion, 2020.（SCI一区）",
        "Dwivedi Y K, Kshetri N, Hughes L, et al. Opinion Paper: So what if ChatGPT wrote it? International Journal of Information Management, 2023.（SSCI一区）",
        "Ebers M, Hoch V R S, Rosenkranz F, et al. The European Commission's Proposal for an Artificial Intelligence Act: A Critical Assessment by Members of the Robotics and AI Law Society. J, 2021.",
        "Fjeld J, Achten N, Hilligoss H, et al. Principled Artificial Intelligence: Mapping Consensus in Ethical and Rights-Based Approaches to Principles for AI. Berkman Klein Center Research Publication, 2020.",
        "Mökander J, Schuett J, Kirk H R, et al. Auditing large language models: A three-layered approach. AI and Ethics, 2023.",
        "Floridi L. The Ethics of Artificial Intelligence: Principles, Challenges, and Opportunities. Oxford University Press, 2023.",
        "国家互联网信息办公室等. 生成式人工智能服务管理暂行办法. 2023.",
        "全国网络安全标准化技术委员会. 人工智能安全治理框架1.0版. 2024.",
        "European Parliament and Council. Regulation (EU) 2024/1689: Artificial Intelligence Act. 2024.",
        "U.S. Government Accountability Office. Artificial Intelligence: Use and Oversight in Financial Services. GAO-25-107197, 2025.",
        "International Monetary Fund. Regulatory Considerations Regarding Accelerated Use of AI in Securities Markets. Technical Notes and Manuals 2025/016, 2025.",
    ]
    recent_count = sum(any(y in ref for y in ["2023", "2024", "2025", "2026"]) for ref in refs)
    add_para(doc, f"注：以下列出 {len(refs)} 条参考文献，其中近三年文献/政策 {recent_count} 条，占比超过一半；正式提交前建议按学校最新版 CCF 与中科院分区目录复核标注。")
    for idx, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"[{idx}] {ref}")
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
