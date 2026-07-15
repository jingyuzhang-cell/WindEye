from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt, Inches


ROOT = Path(__file__).resolve().parents[1]
INPUT_MD = ROOT / "docs" / "第二章_系统功能详细设计_代码重写版.md"
OUTPUT_DOCX = ROOT / "docs" / "第二章_系统功能详细设计_代码重写版.docx"


def set_run_font(run, name: str, size: int | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw in lines:
        if re.fullmatch(r"\|\s*[-: ]+\|\s*", raw) or set(raw.replace("|", "").strip()) <= {"-", ":", " "}:
            continue
        cells = [cell.strip() for cell in raw.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def apply_inline_code(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 10)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Calibri", 11)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = [Inches(1.6)] + [Inches(4.9 / max(1, col_count - 1))] * (col_count - 1)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, "Calibri", 10, bold=(r_idx == 0))
                p.paragraph_format.space_after = Pt(2)
        for c_idx, width in enumerate(widths[:col_count]):
            table.cell(r_idx, c_idx).width = width


def build_doc() -> None:
    text = INPUT_MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_styles(doc)

    code_block: list[str] = []
    in_code = False
    table_block: list[str] = []
    para_buffer: list[str] = []

    def flush_paragraph():
        nonlocal para_buffer
        if not para_buffer:
            return
        text_line = " ".join(item.strip() for item in para_buffer if item.strip())
        if not text_line:
            para_buffer = []
            return
        if re.match(r"^\d+\.\s+", text_line):
            p = doc.add_paragraph(style="List Number")
            apply_inline_code(p, re.sub(r"^\d+\.\s+", "", text_line))
        elif text_line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            apply_inline_code(p, text_line[2:].strip())
        else:
            p = doc.add_paragraph()
            apply_inline_code(p, text_line)
        para_buffer = []

    def flush_table():
        nonlocal table_block
        if table_block:
            add_table(doc, parse_table(table_block))
            table_block = []

    for line in lines:
        if line.startswith("```"):
            flush_paragraph()
            flush_table()
            if in_code:
                p = doc.add_paragraph()
                for idx, code_line in enumerate(code_block):
                    run = p.add_run(code_line)
                    set_run_font(run, "Consolas", 9)
                    if idx != len(code_block) - 1:
                        run.add_break()
                code_block = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_block.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            flush_paragraph()
            table_block.append(line)
            continue
        elif table_block:
            flush_table()

        if not line.strip():
            flush_paragraph()
            continue

        if line.startswith("# "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 1")
            apply_inline_code(p, line[2:].strip())
            continue
        if line.startswith("## "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 1")
            apply_inline_code(p, line[3:].strip())
            continue
        if line.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 2")
            apply_inline_code(p, line[4:].strip())
            continue
        if line.startswith("#### "):
            flush_paragraph()
            p = doc.add_paragraph(style="Heading 3")
            apply_inline_code(p, line[5:].strip())
            continue

        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            flush_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            set_run_font(run, "Calibri", 11, True)
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            apply_inline_code(p, re.sub(r"^\d+\.\s+", "", line.strip()))
            continue

        if line.strip().startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            apply_inline_code(p, line.strip()[2:].strip())
            continue

        para_buffer.append(line)

    flush_paragraph()
    flush_table()

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_doc()
