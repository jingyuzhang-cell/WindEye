# -*- coding: utf-8 -*-
"""Generate a test risk event Word document for upload/QA testing.

The content is aligned with the existing Huachuang knowledge graph sample data,
covering all 7 companies, 5 persons, and multiple event types across the 4-layer KG.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

LQ = '“'  # left Chinese quote
RQ = '”'  # right Chinese quote
ARROW = '→'  # right arrow


def set_cell_font(cell, text, bold=False, size=10.5):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run.bold = bold


def q(text):
    """Wrap text in Chinese quotation marks."""
    return LQ + text + RQ


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # =================================================================
    # TITLE
    # =================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("关于华创系企业风险事件的综合分析报告")
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("报告日期：2025年5月28日　　　密级：内部机密")
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # =================================================================
    # 一、报告摘要
    # =================================================================
    h1 = doc.add_paragraph()
    run = h1.add_run("一、报告摘要")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    p1 = (
        "本报告基于公开市场信息与监管披露文件，对华创控股集团有限公司（以下简称"
        + q("华创集团") + "）及其关联企业（合称" + q("华创系")
        + "）2024年以来的重大风险事件进行综合分析。华创系涉及企业主体包括华创控股集团"
        "有限公司、华创地产股份有限公司、华创贸易有限责任公司、鑫达投资管理有限公司、海通金融服务"
        "有限公司、中远建设工程有限公司及天元科技发展有限公司共七家企业，实际控制人为张明远。"
    )
    doc.add_paragraph(p1)

    p2 = (
        "经梳理，华创系自2024年3月起先后触发五项重大风险事件，涵盖司法执行、刑事立案、股权冻结、"
        "行政调查及债务违约五大类型，事件呈明显因果传导链条。截至报告日，华创系整体风险敞口超过"
        "30亿元，涉及金融机构债权人12家，已引起中国证监会、公安部及最高人民法院的持续关注。"
    )
    doc.add_paragraph(p2)

    # =================================================================
    # 二、涉事主体概况
    # =================================================================
    h2 = doc.add_paragraph()
    run = h2.add_run("二、涉事主体概况")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph("（一）核心企业")

    table = doc.add_table(rows=8, cols=5, style="Light Grid Accent 1")
    headers = ["企业名称", "注册资本（万元）", "经营状态", "风险预警数", "备注"]
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, bold=True, size=10)

    companies = [
        ("华创控股集团有限公司", "500,000", "存续", "12", "核心控股平台，张明远实际控制"),
        ("华创地产股份有限公司", "200,000", "存续", "8", "主要经营实体，债券发行人"),
        ("华创贸易有限责任公司", "50,000", "存续", "3", "供应链平台，李建国担任法人代表"),
        ("鑫达投资管理有限公司", "10,000", "存续", "5", "投资平台，交叉持股关键节点"),
        ("海通金融服务有限公司", "80,000", "吊销", "15", "非法集资涉案主体"),
        ("中远建设工程有限公司", "30,000", "存续", "2", "华创地产控股子公司"),
        ("天元科技发展有限公司", "15,000", "存续", "1", "相对独立，风险较低"),
    ]
    for i, (name, capital, status, warnings, note) in enumerate(companies):
        row = table.rows[i + 1]
        set_cell_font(row.cells[0], name, size=10)
        set_cell_font(row.cells[1], capital, size=10)
        set_cell_font(row.cells[2], status, size=10)
        set_cell_font(row.cells[3], warnings, size=10)
        set_cell_font(row.cells[4], note, size=10)

    doc.add_paragraph()
    doc.add_paragraph("（二）关键自然人")

    person_table = doc.add_table(rows=6, cols=4, style="Light Grid Accent 1")
    person_headers = ["姓名", "职务/角色", "关联企业", "风险等级"]
    for i, h in enumerate(person_headers):
        set_cell_font(person_table.rows[0].cells[i], h, bold=True, size=10)

    persons = [
        ("张明远", "实际控制人/董事长", "华创控股、华创地产、鑫达投资", "高风险"),
        ("李建国", "总经理/法人代表", "华创控股、华创贸易", "中风险"),
        ("王丽华", "财务总监/财务负责人", "华创控股、海通金融", "高风险"),
        ("赵志强", "法人代表/项目经理", "海通金融、中远建设", "高风险"),
        ("陈晓峰", "监事", "鑫达投资、海通金融", "中风险"),
    ]
    for i, (name, role, co_list, risk) in enumerate(persons):
        row = person_table.rows[i + 1]
        set_cell_font(row.cells[0], name, size=10)
        set_cell_font(row.cells[1], role, size=10)
        set_cell_font(row.cells[2], co_list, size=10)
        set_cell_font(row.cells[3], risk, size=10)

    # =================================================================
    # 三、风险事件详情
    # =================================================================
    doc.add_paragraph()
    h3 = doc.add_paragraph()
    run = h3.add_run("三、风险事件详情")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    # --- Event 1 ---
    doc.add_paragraph()
    e1 = doc.add_paragraph()
    run = e1.add_run("事件一：华创集团被列入被执行人（执行案号：(2024)京01执字第01234号）")
    run.font.size = Pt(12)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "发生日期：2024年3月15日\n"
        "事件类型：司法执行\n"
        "影响等级：高风险\n"
        "涉事主体：华创控股集团有限公司\n"
        "执行法院：北京市第一中级人民法院\n"
        "执行标的：3.2亿元人民币\n"
        "事件概述：华创控股集团有限公司因未履行对供应商北京恒通建材有限公司的付款义务，"
        "被北京市第一中级人民法院列为被执行人。经查明，该笔应付账款涉及华创地产股份有限公司"
        "开发项目的建材采购，因海通金融服务有限公司资金链紧张导致集团内部流动性枯竭，"
        "未能按期支付。\n"
        "法律依据：根据《中华人民共和国民事诉讼法》第二百四十二条、《中华人民共和国公司法》"
        "第二十条关于股东不得滥用权利损害公司利益的规定。"
    )

    # --- Event 2 ---
    doc.add_paragraph()
    e2 = doc.add_paragraph()
    run = e2.add_run("事件二：海通金融涉嫌非法集资（案件编号：公经侦(2024)第05678号）")
    run.font.size = Pt(12)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "发生日期：2024年6月20日\n"
        "事件类型：刑事立案\n"
        "影响等级：高风险\n"
        "涉事主体：海通金融服务有限公司、赵志强（法人代表）、王丽华（财务负责人）\n"
        "办案机关：上海市公安局经济犯罪侦查总队\n"
        "涉案金额：超过15亿元人民币\n"
        "涉及投资人：约3,200名自然人投资者\n"
        "事件概述：海通金融服务有限公司因涉嫌非法吸收公众存款，被上海市公安局经侦总队立案"
        "侦查。初步查明，该公司自2022年起以" + q("私募股权投资基金") + "名义，通过鑫达投资管理有限公司"
        "募集资金，承诺年化收益率8%-15%，资金实际流入华创系体内用于房地产项目开发。该行为"
        "涉嫌违反《中华人民共和国刑法》第一百七十六条关于非法吸收公众存款罪的规定。\n"
        "关联影响：该事件触发鑫达投资持有的华创地产10%股权被冻结（见事件三），并引发"
        "监管层对华创系整体合规性的关注（见事件四）。"
    )

    # --- Event 3 ---
    doc.add_paragraph()
    e3 = doc.add_paragraph()
    run = e3.add_run("事件三：鑫达投资股权冻结（裁定书编号：(2024)沪74民初字第00345号）")
    run.font.size = Pt(12)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "发生日期：2024年9月10日\n"
        "事件类型：股权冻结\n"
        "影响等级：中风险\n"
        "涉事主体：鑫达投资管理有限公司、华创地产股份有限公司\n"
        "执行法院：上海市金融法院\n"
        "冻结标的：鑫达投资持有的华创地产10%股权（对应出资额20,000万元）\n"
        "事件概述：因海通金融非法集资案件侦查需要，上海市金融法院依法冻结鑫达投资管理有限公司"
        "持有的华创地产股份有限公司10%股权。该冻结直接影响华创地产的公司治理结构及后续融资计划。"
        "值得注意的是，该笔交叉持股正是华创系股权穿透异常的核心环节——华创贸易持有鑫达投资15%"
        "股权，鑫达投资持有华创地产10%股权，形成隐蔽的交叉持股闭环。"
    )

    # --- Event 4 ---
    doc.add_paragraph()
    e4 = doc.add_paragraph()
    run = e4.add_run("事件四：证监会立案调查华创系（调查通知编号：证监调查字(2024)第00128号）")
    run.font.size = Pt(12)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "发生日期：2024年12月5日\n"
        "事件类型：行政调查\n"
        "影响等级：高风险\n"
        "涉事主体：华创控股集团有限公司、华创地产股份有限公司、海通金融服务有限公司\n"
        "调查机构：中国证券监督管理委员会\n"
        "调查事项：（1）华创控股未及时披露对海通金融的5亿元担保事项；"
        "（2）华创系内部关联交易金额占营收比例超60%，涉嫌利益输送；"
        "（3）海通金融非法集资案中的信息披露违规行为。\n"
        "事件概述：因海通金融非法集资案及市场舆情发酵，中国证监会对华创系三家公司同步启动"
        "立案调查。调查组重点关注华创系的信息披露合规性、关联交易公允性及公司治理有效性。"
        "张明远作为实际控制人，被要求限期说明华创系整体资金往来情况。\n"
        "法律依据：《中华人民共和国证券法》第八十二条关于信息披露义务的规定、"
        "《中华人民共和国公司法》第二十条关于股东禁止滥用权利的规定。"
    )

    # --- Event 5 ---
    doc.add_paragraph()
    e5 = doc.add_paragraph()
    run = e5.add_run("事件五：华创地产债券违约（债券代码：20华创01，ISIN：CN01200001234）")
    run.font.size = Pt(12)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "发生日期：2025年1月18日\n"
        "事件类型：债务违约\n"
        "影响等级：高风险\n"
        "涉事主体：华创地产股份有限公司\n"
        "主承销商：国泰君安证券股份有限公司\n"
        "受托管理人：中信信托有限责任公司\n"
        "违约金额：本息合计8.5亿元人民币\n"
        "事件概述：华创地产股份有限公司发行的" + q("20华创01") + "公司债券未能按期兑付本息，构成"
        "实质性违约。该债券发行规模10亿元，票面利率6.5%，期限3+2年。违约直接原因系华创集团"
        "被列入被执行人后，主要银行账户被冻结，流动性枯竭；深层原因为海通金融非法集资案导致"
        "集团融资渠道全面受阻。\n"
        "关联影响：2025年2月1日，中诚信国际信用评级有限责任公司将华创地产主体信用等级由AA-"
        "下调至BBB，展望负面。"
    )

    # =================================================================
    # 四、股权穿透与关联分析
    # =================================================================
    doc.add_paragraph()
    h4 = doc.add_paragraph()
    run = h4.add_run("四、股权穿透与关联分析")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "（一）股权层级结构\n"
        "华创控股集团有限公司作为顶层控股平台，持有华创地产70%股权、华创贸易51%股权、"
        "鑫达投资80%股权。华创地产进一步持有中远建设60%股权。股权层级最深达到3级，"
        "整体架构呈现" + q("金字塔+交叉持股") + "的复杂特征。"
    )
    doc.add_paragraph(
        "（二）交叉持股闭环\n"
        "华创贸易有限责任公司出资7,500万元持有鑫达投资管理有限公司15%股权，而鑫达投资又出资"
        "20,000万元持有华创地产股份有限公司10%股权，华创地产由华创控股控股，华创贸易亦由华创"
        "控股控股51%。该交叉持股结构在法律形式上虽未直接违反现行法规，但在实质上削弱了公司"
        "治理的透明度，增加了风险传导的隐蔽性。该交叉持股闭环已被识别为" + q("股权穿透异常") + "风险特征"
        "（特征编号F001）。"
    )
    doc.add_paragraph(
        "（三）连环担保网络\n"
        "华创地产为海通金融提供5亿元担保（2024年3月15日起），海通金融又为华创贸易提供1亿元"
        "担保（2024年6月1日起）。该连环担保形成的或有负债总额达6亿元，占华创地产净资产比例"
        "超过50%，已被识别为" + q("担保链风险") + "特征（特征编号F002）。一旦任一节点发生违约，担保链"
        "将产生多米诺骨牌效应。"
    )
    doc.add_paragraph(
        "（四）高管交叉任职\n"
        "张明远同时担任华创控股实际控制人、华创地产董事长及鑫达投资执行董事，王丽华同时担任"
        "华创控股财务总监和海通金融财务负责人。该高管交叉任职模式违反公司治理的最佳实践，"
        "已被识别为" + q("关键人兼职") + "风险特征（特征编号F003）。"
    )

    # =================================================================
    # 五、风险特征识别
    # =================================================================
    doc.add_paragraph()
    h5 = doc.add_paragraph()
    run = h5.add_run("五、风险特征识别")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph("基于对上述风险事件的综合分析，识别出以下五项核心风险特征：")

    risk_features = [
        ("F001 股权穿透异常 -- 交叉持股环路",
         "华创贸易持有鑫达投资15%股权，鑫达投资持有华创地产10%股权，华创地产为华创控股子公司，"
         "形成隐蔽的交叉持股闭环，股权层级超过3级。"
         "触发事件：华创集团被列入被执行人、鑫达投资股权冻结。"
         "风险因子：股权层级>3级（因子编号FK001）、司法被执行>=1次（因子编号FK005）。"),
        ("F002 担保链风险 -- 连环担保",
         "华创地产为海通金融提供5亿元担保，海通金融为华创贸易提供1亿元担保，形成集团内连环"
         "担保网络，担保总额/净资产比率超过50%。"
         "触发事件：海通金融涉嫌非法集资。"
         "风险因子：担保金额/净资产>50%（因子编号FK002）。"),
        ("F003 高管交叉任职 -- 关键人兼职",
         "实际控制人张明远同时在华创控股、华创地产、鑫达投资担任董事/执行董事，财务总监王丽华"
         "同时服务华创控股和海通金融，高管兼任企业数>=3家。"
         "触发事件：证监会立案调查华创系。"
         "风险因子：高管兼任>=3家（因子编号FK003）。"),
        ("F004 资金异常流动 -- 关联交易集中",
         "2024年度华创系内部关联交易额占营收比例超过60%，显著高于行业平均水平（约15%），"
         "可能存在利益输送和资金占用嫌疑。"
         "触发事件：证监会立案调查华创系。"
         "风险因子：关联交易占比>50%（因子编号FK004）。"),
        ("F005 舆情负面 -- 媒体负面报道密集",
         "近6个月华创系相关负面舆情报道37篇，涉及债务违约、股权冻结、监管调查、非法集资等"
         "关键词。"
         "触发事件：华创地产债券违约。"
         "风险因子：债券违约（因子编号FK006）。"),
    ]

    for title_text, body_text in risk_features:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.font.size = Pt(11)
        run.font.name = "黑体"
        run.bold = True
        doc.add_paragraph(body_text)

    # =================================================================
    # 六、涉及法规分析
    # =================================================================
    doc.add_paragraph()
    h6 = doc.add_paragraph()
    run = h6.add_run("六、涉及法规分析")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    regulation_mapping = [
        ("《中华人民共和国公司法》第二十条",
         "公司股东应当遵守法律、行政法规和公司章程，依法行使股东权利，不得滥用股东权利损害"
         "公司或者其他股东的利益。适用于华创控股作为控股股东对子公司进行的关联交易和担保行为。"
         "适用事件：华创集团被列入被执行人、证监会立案调查华创系。"),
        ("《中华人民共和国证券法》第八十二条",
         "上市公司董事、监事、高级管理人员应当保证上市公司所披露的信息真实、准确、完整。"
         "适用于华创系未及时披露对外担保和关联交易的信息披露违规行为。"
         "适用事件：证监会立案调查华创系、华创控股信披违规处罚。"),
        ("《中华人民共和国刑法》第一百七十六条",
         "非法吸收公众存款或者变相吸收公众存款，扰乱金融秩序的，处三年以下有期徒刑或者拘役。"
         "适用于海通金融以私募基金名义向社会公众募集资金的行为。"
         "适用事件：海通金融涉嫌非法集资。"),
        ("《最高人民法院关于适用<中华人民共和国公司法>若干问题的规定（五）》",
         "进一步明确关联交易损害的认定标准和股东代表诉讼程序。"
         "适用于华创系内部关联交易的合规性审查。"
         "适用事件：华创贸易关联方资金占用。"),
        ("《中华人民共和国民法典》",
         "担保合同的成立、效力及担保人责任认定。"
         "适用于华创系连环担保网络中各项担保行为的法律效力认定。"
         "适用事件：华创地产为海通金融提供担保。"),
    ]

    for title_text, body_text in regulation_mapping:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.font.size = Pt(11)
        run.font.name = "黑体"
        run.bold = True
        doc.add_paragraph(body_text)

    # =================================================================
    # 七、风险传导路径
    # =================================================================
    doc.add_paragraph()
    h7 = doc.add_paragraph()
    run = h7.add_run("七、风险传导路径")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "华创系风险事件的传导呈现清晰的因果链条，可分为两条主要传导路径："
    )
    doc.add_paragraph(
        "路径一（司法 " + ARROW + " 刑事 " + ARROW + " 监管 " + ARROW + " 信用）：\n"
        "华创集团被列入被执行人（2024.03.15）" + ARROW + " 子公司银行账户冻结、流动枯竭 " + ARROW + " 华创地产"
        "债券违约（2025.01.18）" + ARROW + " 华创地产信用评级下调（2025.02.01）。\n"
        "该路径展示了司法执行如何通过资金链传导最终演变为信用风险事件。"
    )
    doc.add_paragraph(
        "路径二（刑事 " + ARROW + " 司法 " + ARROW + " 监管）：\n"
        "海通金融非法集资案发（2024.06.20）" + ARROW + " 鑫达投资股权冻结（2024.09.10）" + ARROW + " 证监会立案"
        "调查华创系三家公司（2024.12.05）" + ARROW + " 华创控股信披违规处罚（2025.03.22）。\n"
        "该路径展示了刑事案件如何触发监管调查、进而引发行政处罚的连锁反应。"
    )

    chain = (
        "以上两条路径最终汇合，共同作用于整个华创系，形成"
        + q("司法执行 " + ARROW + " 刑事立案 " + ARROW + " 股权冻结 "
           + ARROW + " 行政调查 " + ARROW + " 债务违约 " + ARROW + " 评级下调 "
           + ARROW + " 行政处罚")
        + "的完整风险传导链条。"
    )
    doc.add_paragraph(chain)

    # =================================================================
    # 八、结论与建议
    # =================================================================
    doc.add_paragraph()
    h8 = doc.add_paragraph()
    run = h8.add_run("八、结论与建议")
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run.bold = True

    doc.add_paragraph(
        "（一）综合结论\n"
        "华创系风险事件具有成因复杂、传导迅速、影响广泛三个突出特征。从2024年3月首个风险事件"
        "爆出，到2025年4月多个事件并发，仅13个月时间即形成了覆盖全部七家企业的系统性风险。"
        "海通金融非法集资案是整个风险链的关键引爆点，而华创系自身的股权穿透异常、连环担保网络"
        "及高管交叉任职等结构性问题，构成了风险快速传导的制度基础。"
    )
    doc.add_paragraph(
        "（二）风险处置建议\n"
        "1. 建议对张明远、王丽华、赵志强三人实施限制消费措施，防止转移资产；\n"
        "2. 建议上海金融法院对鑫达投资所持华创地产股权进行司法拍卖，以清偿相关债务；\n"
        "3. 建议证监会向华创控股、华创地产发送监管问询函，要求限期补充披露关联交易和对外担保事项；\n"
        "4. 建议公安机关移送海通金融非法集资案相关材料至检察院审查起诉；\n"
        "5. 建议将华创系全部七家企业纳入监管重点监控名单，持续跟踪风险演变。"
    )
    doc.add_paragraph(
        "（三）后续关注事项\n"
        "1. 华创地产" + q("20华创01") + "债券持有人会议进展及兑付方案；\n"
        "2. 海通金融非法集资案的侦查进展及受害人退赔情况；\n"
        "3. 证监会现场检查结果及可能的行政处罚决定；\n"
        "4. 华创系是否触发《中华人民共和国企业破产法》第二条规定的破产原因；\n"
        "5. 天元科技发展有限公司作为风险相对较低的主体，是否存在被风险传导的可能。"
    )

    # =================================================================
    # Footer
    # =================================================================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("报告编制：风控合规部\n审核：风险管理委员会\n报告编号：RC-2025-0528-001")
    run.font.size = Pt(9)
    run.font.name = "宋体"
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = "d:/Code/WindEye/tests/data/华创系风险事件综合分析报告.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path):,} bytes")
