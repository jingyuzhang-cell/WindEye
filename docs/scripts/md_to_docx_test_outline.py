"""Convert v1.0 系统测试大纲.md to .docx using python-docx."""

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os


# ── Helpers ──
def set_cell_font(cell, text, bold=False, size=9, font_name='Microsoft YaHei'):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.bold = bold
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h


def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)


def shade_header_row(table, color='1F4E79'):
    for cell in table.rows[0].cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True


# ── Read .md file ──
md_path = r'D:\Code\WindEye\docs\formal\v1.0 系统测试大纲.md'
with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')

# ── Create document ──
doc = Document()

# Page setup (A4)
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── Parse and render markdown ──
i = 0
in_code = False
table_rows = []
skip_separators = True

while i < len(lines):
    line = lines[i]

    # Skip YAML frontmatter
    if line.strip() == '---' and i < 5:
        i += 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1
        continue

    # Code blocks
    if line.strip().startswith('```'):
        in_code = not in_code
        i += 1
        continue
    if in_code:
        i += 1
        continue

    # Table handling
    if line.strip().startswith('|') and line.strip().endswith('|'):
        if '---' in line and skip_separators:
            skip_separators = False
            table_rows.append(line)
            i += 1
            continue

        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        table_rows.append(cells)

        # Check if next line continues the table
        if i + 1 >= len(lines) or not (
            lines[i + 1].strip().startswith('|')
            and lines[i + 1].strip().endswith('|')
        ):
            if len(table_rows) >= 2:
                # Remove separator row
                data_rows = [
                    r for r in table_rows if not all('---' in c or c == '---' for c in r)
                ]
                if data_rows:
                    num_cols = max(len(r) for r in data_rows)
                    table = doc.add_table(rows=len(data_rows), cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(table)

                    for ri, row_data in enumerate(data_rows):
                        for ci, cell_text in enumerate(row_data):
                            if ci < num_cols:
                                is_header = ri == 0
                                set_cell_font(
                                    table.rows[ri].cells[ci],
                                    cell_text,
                                    bold=is_header,
                                    size=8 if not is_header else 9,
                                )

                    if data_rows:
                        shade_header_row(table)
                    doc.add_paragraph()  # spacing after table
            table_rows = []
            skip_separators = True
        i += 1
        continue

    # Reset skip flag when not in table
    skip_separators = True

    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(22)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        doc.add_paragraph()
    elif line.startswith('## '):
        add_heading_styled(doc, line[3:].strip(), level=2)
    elif line.startswith('### '):
        add_heading_styled(doc, line[4:].strip(), level=3)
    elif line.startswith('#### '):
        add_heading_styled(doc, line[5:].strip(), level=4)
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.strip():
                run = p.add_run(part)
            else:
                continue
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    elif line.strip():
        text = line.strip()
        if '**' in text:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part:
                    run = p.add_run(part)
                else:
                    continue
                run.font.size = Pt(10.5)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        else:
            add_body(doc, text)

    i += 1

# ── Save ──
output_path = r'D:\Code\WindEye\docs\formal\v1.0 系统测试大纲.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
print(f'File size: {os.path.getsize(output_path)} bytes')
print('Done!')
