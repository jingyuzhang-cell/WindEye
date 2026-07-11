from __future__ import annotations

import urllib.parse
import urllib.request
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from api_test_common import (
    OUTPUT_DIR,
    REPORT_OUTPUT_DIR,
    call_json,
    parse_args,
    print_row,
    resolve_company_node,
    save_result,
    wrap_result,
)


def inspect_docx(path: str, company: str) -> dict:
    docx_path = REPORT_OUTPUT_DIR / path if not path or ":" not in path else None
    if docx_path is None:
        from pathlib import Path

        docx_path = Path(path)
    exists = docx_path.exists()
    size = docx_path.stat().st_size if exists else 0
    readable = False
    contains_company = False
    paragraph_count = 0
    table_count = 0
    title = ""
    error = ""
    if exists:
        try:
            from docx import Document

            doc = Document(str(docx_path))
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            text = "\n".join(p.text for p in doc.paragraphs)
            title = doc.core_properties.title or ""
            readable = True
            contains_company = company in text
        except Exception as exc:
            try:
                text = extract_docx_text(docx_path)
                readable = True
                contains_company = company in text
                paragraph_count = max(1, text.count("\n") + 1) if text else 0
                title = "WindEye 协同治理报告" if "WindEye 协同治理报告" in text else ""
                error = f"python-docx unavailable, zip fallback used: {exc}"
            except Exception as fallback_exc:
                error = f"{exc}; zip fallback failed: {fallback_exc}"
    return {
        "path": str(docx_path),
        "exists": exists,
        "sizeBytes": size,
        "readable": readable,
        "containsCompany": contains_company,
        "paragraphCount": paragraph_count,
        "tableCount": table_count,
        "title": title,
        "error": error,
    }


def extract_docx_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return "\n".join(t.text for t in root.iter(ns + "t") if t.text)


def download_report_file(base_url: str, download_url: str, filename: str) -> dict:
    if not download_url or not filename:
        return {
            "downloaded": False,
            "path": "",
            "sizeBytes": 0,
            "error": "missing downloadUrl or fileName",
        }
    url = urllib.parse.urljoin(base_url.rstrip("/") + "/", download_url.lstrip("/"))
    output_path = OUTPUT_DIR / filename
    try:
        with urllib.request.urlopen(url, timeout=120) as resp:
            data = resp.read()
        output_path.write_bytes(data)
        return {
            "downloaded": True,
            "url": url,
            "path": str(output_path),
            "sizeBytes": output_path.stat().st_size,
            "error": "",
        }
    except Exception as exc:
        return {
            "downloaded": False,
            "url": url,
            "path": str(output_path),
            "sizeBytes": 0,
            "error": str(exc),
        }


def main() -> int:
    base_url, company = parse_args()
    resolved = resolve_company_node(base_url, company)
    seed_name = resolved.get("nodeName") or company
    seed_id = resolved.get("nodeId", "")
    api = "/api/v1/governance/compliance-report"
    body = {
        "query": f"分析{seed_name}的风险传导、群体发现和协同治理社区报告",
        "seedNames": [],
        "seedIds": [seed_id] if seed_id else [],
        "focusEntities": [seed_name],
        "maxHop": 2,
        "maxPathLength": 4,
        "method": "auto",
        "communityMode": "expanded",
        "minCommunitySize": 2,
        "pathLimit": 5000,
        "maxNodes": 1000,
        "riskPathLimit": 10,
        "maxBranchPerNode": 10,
        "minRiskScore": 0,
        "includeRawSubgraph": True,
        "includeCommunityGraph": True,
        "includeCommunityPath": True,
        "includeNodePath": True,
        "exportFormats": ["docx"],
        "reportOptions": {
            "delivery": "metadata",
            "includeDownloadUrl": True,
            "includeServerPath": False,
            "includeBase64": False,
        },
        "sessionId": "xugong-api-test",
        "roundId": 1,
    }
    result = call_json(base_url, "POST", api, body)
    response = result.get("response")
    report_payload = response if isinstance(response, dict) else {}
    compliance = report_payload.get("compliance", {}) if isinstance(report_payload.get("compliance"), dict) else {}
    indicators = (
        report_payload.get("complianceIndicators", {})
        if isinstance(report_payload.get("complianceIndicators"), dict)
        else {}
    )
    governance = report_payload.get("governance", {}) if isinstance(report_payload.get("governance"), dict) else {}
    report = report_payload.get("report", {}) if isinstance(report_payload.get("report"), dict) else {}
    view_model = report_payload.get("viewModel", {}) if isinstance(report_payload.get("viewModel"), dict) else {}
    pipeline_trace = report_payload.get("pipelineTrace", {}) if isinstance(report_payload.get("pipelineTrace"), dict) else {}
    export_files = report_payload.get("exportFiles", {}) if isinstance(report_payload.get("exportFiles"), dict) else {}
    docx_meta = export_files.get("docx", {}) if isinstance(export_files.get("docx"), dict) else {}
    download_result = download_report_file(
        base_url,
        str(docx_meta.get("downloadUrl") or ""),
        str(docx_meta.get("fileName") or ""),
    )
    docx_inspection = inspect_docx(download_result.get("path") or str(docx_meta.get("filePath") or ""), seed_name)
    report_word_complete = bool(
        result.get("statusCode") == 200
        and result.get("isJson")
        and report_payload.get("success") is True
        and report_payload.get("defaultFormat") == "docx"
        and isinstance(compliance, dict)
        and isinstance(indicators, dict)
        and isinstance(governance.get("actions", []), list)
        and isinstance(report.get("markdownReport", ""), str)
        and report.get("format") == "docx"
        and str(report.get("fileName", "")).endswith(".docx")
        and isinstance(docx_meta.get("downloadUrl"), str)
        and docx_meta.get("downloadUrl")
        and not docx_meta.get("filePath")
        and download_result.get("downloaded")
        and docx_inspection["exists"]
        and docx_inspection["sizeBytes"] > 0
        and docx_inspection["readable"]
        and docx_inspection["containsCompany"]
        and isinstance(view_model.get("highlightNodeIds", []), list)
        and pipeline_trace.get("communityDiscoveryGenerated") is True
        and pipeline_trace.get("riskPathsGenerated") is True
    )
    payload = wrap_result(
        api,
        result,
        expected_status=200,
        extra={
            "company": company,
            "seedName": seed_name,
            "seedId": seed_id,
            "resolution": resolved,
            "reportWordComplete": report_word_complete,
            "defaultFormat": report_payload.get("defaultFormat"),
            "exportFiles": export_files,
            "downloadResult": download_result,
            "docxInspection": docx_inspection,
            "subject": report_payload.get("subject"),
            "riskLevel": compliance.get("riskLevel"),
            "complianceScore": compliance.get("score"),
            "indicatorScore": indicators.get("totalScore"),
            "pipelineTrace": pipeline_trace,
            "internalCommunityGenerated": pipeline_trace.get("communityDiscoveryGenerated"),
            "internalRiskPathsGenerated": pipeline_trace.get("riskPathsGenerated"),
            "internalCommunityCount": pipeline_trace.get("communityCount"),
            "internalRiskPathCount": pipeline_trace.get("riskPathCount"),
            "governanceActionCount": len(governance.get("actions", [])) if isinstance(governance, dict) else 0,
            "recommendationCount": len(report.get("recommendations", [])) if isinstance(report, dict) else 0,
            "markdownLength": len(report.get("markdownReport", "")) if isinstance(report, dict) else 0,
            "highlightNodeCount": len(view_model.get("highlightNodeIds", [])) if isinstance(view_model, dict) else 0,
            "highlightEdgeCount": len(view_model.get("highlightEdgeIds", [])) if isinstance(view_model, dict) else 0,
            "communityReportSourceCount": len(report_payload.get("communityReportSources", []))
            if isinstance(report_payload.get("communityReportSources"), list)
            else 0,
        },
    )
    save_result("05_compliance_report_xugong.json", payload)
    if report_payload:
        save_result("05_compliance_report_xugong_response.json", report_payload)
    print_row(payload)
    return 0 if report_word_complete else 1


if __name__ == "__main__":
    raise SystemExit(main())
