"""API route definitions for BiDA-KG backend."""

import json as _json
import logging
from dataclasses import asdict
from uuid import uuid4

from fastapi import APIRouter, Request, File, UploadFile
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel, Field
from typing import Any, Literal

from config.settings import settings
from core.exceptions import BiDAError
from core.models import ApiSuccessResponse, TraceContext
from core.tracing import get_trace_id

router = APIRouter()
_logger = logging.getLogger("api.router")


# ── Risk Analysis request models ──────────────────────────────────

class RiskAnalyzeRequest(BaseModel):
    query: str = Field(..., description="User query or trigger description.")
    focus_entities: list[str] = Field(default_factory=list, description="Entity names to focus on.")
    max_hop: int = Field(default=3, ge=1, le=5, description="Max analysis hop count.")
    trigger_event: str | None = Field(default=None, description="Trigger event type.")


class ChatRecommendRequest(BaseModel):
    query: str = Field(..., description="Current user query.")
    history: list[str] = Field(default_factory=list, description="Conversation history.")
    session_id: str = Field(default_factory=lambda: f"sess-{uuid4().hex[:10]}", alias="sessionId")
    round_id: int = Field(default=1, alias="roundId")


class RouteRequest(BaseModel):
    query: str = Field(..., description="User query to route.")


class RouteResponse(BaseModel):
    route: Literal["graph", "analysis", "clarify", "risk"]
    clarify_message: str | None = Field(default=None, description="Clarification question when route=clarify")


class AnalyzeRequest(BaseModel):
    query: str = Field(..., description="Analysis request query.")


class RiskStreamRequest(BaseModel):
    query: str = Field(default="", description="User query for risk analysis.")
    focus_entities: list[str] = Field(default_factory=list, alias="focusEntities", description="Pre-extracted focus entities.")
    file_content: str | None = Field(default=None, alias="fileContent", description="Uploaded file text for entity extraction.")
    session_id: str = Field(default="", alias="sessionId")
    round_id: int = Field(default=1, alias="roundId")
    community_id: int | None = Field(default=None, alias="communityId")
    max_hop: int = Field(default=3, ge=1, le=5, alias="maxHop")


class ReportDocxExportRequest(BaseModel):
    report: dict = Field(default_factory=dict, description="Risk report payload to export.")
    report_id: str = Field(default="", alias="reportId")
    query_text: str = Field(default="", alias="queryText")


class RiskPathsRequest(BaseModel):
    query: str = Field(default="", description="User question or risk-path analysis prompt.")
    focus_entities: list[str] = Field(default_factory=list, alias="focusEntities")
    source_entity: str | None = Field(default=None, alias="sourceEntity")
    target_entities: list[str] = Field(default_factory=list, alias="targetEntities")
    max_hop: int = Field(default=3, ge=1, le=5, alias="maxHop")
    mode: Literal["node", "community", "both"] = "both"
    relation_types: list[str] = Field(default_factory=list, alias="relationTypes")
    subgraph: dict[str, Any] | None = None
    communities: dict[str, Any] | None = None
    entity_community_map: dict[str, Any] | None = Field(default=None, alias="entityCommunityMap")
    file_content: str | None = Field(default=None, alias="fileContent")
    session_id: str = Field(default="", alias="sessionId")
    round_id: int = Field(default=1, alias="roundId")


class GovernanceReportRequest(BaseModel):
    query: str = Field(..., description="User question for collaborative governance report.")
    focus_entities: list[str] = Field(default_factory=list, alias="focusEntities")
    max_hop: int = Field(default=3, ge=1, le=5, alias="maxHop")
    file_content: str | None = Field(default=None, alias="fileContent")
    community_result: dict[str, Any] | None = Field(default=None, alias="communityResult")
    risk_paths: dict[str, Any] | list[dict[str, Any]] | None = Field(default=None, alias="riskPaths")
    compliance_context: dict[str, Any] | None = Field(default=None, alias="complianceContext")
    export_formats: list[Literal["markdown", "docx", "pdf"]] = Field(default_factory=list, alias="exportFormats")
    session_id: str = Field(default="", alias="sessionId")
    round_id: int = Field(default=1, alias="roundId")


class EntityAliasRequest(BaseModel):
    alias: str = Field(..., description="User-entered short name or alias.")
    canonical_name: str = Field(..., alias="canonicalName", description="Confirmed canonical KG entity name.")
    kg_node_id: str = Field(default="", alias="kgNodeId")
    entity_type: str = Field(default="", alias="entityType")
    source: str = Field(default="user_confirmed")


def _api_ok(data: Any, msg: str = "success") -> dict[str, Any]:
    return {"code": 0, "msg": msg, "data": data}


def _api_error(code: int, msg: str, data: Any = None) -> dict[str, Any]:
    return {"code": code, "msg": msg, "data": data}


def _build_risk_path_query(payload: RiskPathsRequest) -> str:
    if payload.query.strip():
        return payload.query.strip()
    entities = payload.focus_entities[:]
    if payload.source_entity:
        entities.insert(0, payload.source_entity)
    if payload.target_entities:
        entities.extend(payload.target_entities)
    unique_entities = []
    for entity in entities:
        if entity and entity not in unique_entities:
            unique_entities.append(entity)
    entity_text = "、".join(unique_entities) if unique_entities else "风险主体"
    mode_text = {
        "node": "具体节点路径",
        "community": "群体之间传导路径",
        "both": "群体之间传导路径和具体节点路径",
    }.get(payload.mode, "风险传导路径")
    return f"分析{entity_text}的风险传导过程，输出{mode_text}"


async def _collect_unified_analysis(
    kg_system: Any,
    *,
    query: str,
    session_id: str,
    round_id: int,
    max_hop: int,
    file_content: str | None = None,
) -> dict[str, Any]:
    from dra_ma.orchestrator.unified_engine import UnifiedEngine

    unified = UnifiedEngine(dra_engine=kg_system, demo=False)
    collected: dict[str, Any] = {
        "events": [],
        "subgraph": None,
        "entity_stats": None,
        "community": None,
        "entity_community_map": None,
        "candidate_risk_paths": [],
        "risk_paths": None,
        "anomaly_findings": [],
        "compliance_matches": [],
        "risk_scores": None,
        "governance_plan": None,
        "report": None,
        "done": None,
    }

    async for line in unified.stream(
        query=query,
        session_id=session_id or f"sess-{uuid4().hex[:10]}",
        round_id=round_id,
        max_hop=max_hop,
        intent_hint="risk_analysis",
        file_content=file_content,
    ):
        event = _json.loads(line) if isinstance(line, str) else line
        event_type = event.get("type", "")
        data = event.get("data", event)
        collected["events"].append({"type": event_type, "status": event.get("status"), "data": data})
        if event_type == "subgraph":
            collected["subgraph"] = data
        elif event_type == "entity_stats":
            collected["entity_stats"] = data
        elif event_type == "community":
            collected["community"] = data
        elif event_type == "entity_community_map":
            collected["entity_community_map"] = data
        elif event_type == "candidate_risk_paths":
            collected["candidate_risk_paths"] = data if isinstance(data, list) else []
        elif event_type == "risk_paths":
            collected["risk_paths"] = data
        elif event_type == "anomaly_findings":
            collected["anomaly_findings"] = data if isinstance(data, list) else []
        elif event_type == "compliance":
            collected["compliance_matches"] = data if isinstance(data, list) else []
        elif event_type == "scoring":
            collected["risk_scores"] = data
        elif event_type == "governance":
            collected["governance_plan"] = data
        elif event_type == "report":
            collected["report"] = data
        elif event_type == "done":
            collected["done"] = data
    return collected


def create_routes(app, kg_system, risk_engine=None):
    """Register all API routes on the FastAPI application."""

    # ═══════════════════════════════════════════════════════════════
    # Module B: 知识图谱展示 — Graph visualization routes
    # ═══════════════════════════════════════════════════════════════
    from api.graph_routes import router as graph_router
    app.include_router(graph_router)

    # ═══════════════════════════════════════════════════════════════
    # Module A: 知识图谱构建 — Pipeline management routes
    # ═══════════════════════════════════════════════════════════════
    from api.pipeline_routes import router as pipeline_router
    app.include_router(pipeline_router)

    # ═══════════════════════════════════════════════════════════════
    # Module C: 风险问答 — Chat / Risk analysis / Ticket routes
    # ═══════════════════════════════════════════════════════════════
    # Governance routes (community discovery, risk paths)
    from api.governance_routes import public_router as public_governance_router
    from api.governance_routes import router as governance_router
    app.include_router(governance_router)
    app.include_router(public_governance_router)

    # ═══════════════════════════════════════════════════════════════
    # [ALL] Authentication & Admin — real auth replacing mock endpoints
    # ═══════════════════════════════════════════════════════════════
    from api.auth_routes import router as auth_router
    app.include_router(auth_router)

    from api.admin_routes import router as admin_router
    app.include_router(admin_router)

    from api.chat_session_routes import router as chat_session_router
    app.include_router(chat_session_router)

    @app.get("/health")
    def health(request: Request) -> dict[str, str]:
        return {"status": "ok", "traceId": getattr(request.state, "trace_id", get_trace_id())}

    # ── Backward-compatible auth endpoints (proxy to real auth) ──

    @app.post("/api/login/account")
    async def login_account_compat(request: Request):
        """Compatibility: proxy POST /api/login/account → POST /api/v1/auth/login."""
        import json as _json

        body = await request.body()
        try:
            body_data = _json.loads(body) if body else {}
        except _json.JSONDecodeError:
            body_data = {}

        username = body_data.get("username", "")
        password = body_data.get("password", "")

        # Fallback to mock login when auth is disabled (MySQL not available)
        if not settings.AUTH_ENABLED:
            if not username or not password:
                return {"status": "error", "type": body_data.get("type", "account"), "currentAuthority": "guest"}
            return {"status": "ok", "type": body_data.get("type", "account"), "currentAuthority": "admin"}

        from api.admin_schemas import LoginRequest
        from api.auth_routes import login as real_login

        req = LoginRequest(
            username=username,
            password=password,
            autoLogin=body_data.get("autoLogin", False),
            type=body_data.get("type"),
        )
        return await real_login(req, request)

    @app.get("/api/currentUser")
    async def current_user_compat(request: Request):
        """Compatibility: proxy GET /api/currentUser → GET /api/v1/auth/me."""
        if not settings.AUTH_ENABLED:
            return {
                "data": {
                    "name": "Admin",
                    "avatar": "https://gw.alipayobjects.com/zos/antfincdn/XAosXuNZyF/BiazfanxmamNRoxxVxka.png",
                    "userid": "00000001",
                    "email": "admin@ant.design",
                    "phone": "138****8888",
                    "access": "admin",
                },
                "success": True,
            }

        from api.auth_routes import current_user as real_current_user
        try:
            return await real_current_user(request)
        except Exception:
            # Fallback if MySQL unavailable
            return {
                "data": {
                    "name": "Admin",
                    "userid": "00000001",
                    "access": "admin",
                },
                "success": True,
            }

    @app.post("/api/login/outLogin")
    async def out_login_compat(request: Request):
        """Compatibility: proxy POST /api/login/outLogin → POST /api/v1/auth/logout."""
        if not settings.AUTH_ENABLED:
            return {"data": {}, "success": True}

        from api.auth_routes import logout as real_logout
        return await real_logout(request)

    @app.get("/api/v1/entities/search")
    async def search_entities(q: str, type: str = "COMPANY", limit: int = 30):
        """Search KG subject-layer candidates for abbreviation confirmation."""
        try:
            from dra_ma.tools.entity_resolver import EntityResolver

            query = (q or "").strip()
            if not query:
                return _api_error(40001, "查询实体不能为空", None)
            safe_limit = max(1, min(int(limit or 10), 30))
            preferred_type = (type or "COMPANY").strip().upper() or None
            resolver = EntityResolver()
            candidates = await resolver.search_candidates(
                query,
                limit=safe_limit * 2,
                preferred_type=preferred_type,
            )
            subject_candidates = [
                item
                for item in candidates
                if "Subject" in item.labels
                or item.entity_type.upper() in {"COMPANY", "PERSON"}
            ][:safe_limit]
            requires_confirmation = (
                len(subject_candidates) != 1
                or subject_candidates[0].match_score < 0.95
                or len(query) <= 12
            )
            return _api_ok({
                "query": query,
                "type": preferred_type,
                "count": len(subject_candidates),
                "requires_confirmation": bool(subject_candidates) and requires_confirmation,
                "auto_resolved": len(subject_candidates) == 1 and not requires_confirmation,
                "candidates": [asdict(item) for item in subject_candidates],
            })
        except Exception as exc:
            _logger.exception("[EntitySearch] failed: %s", exc)
            return _api_error(50003, f"主体候选检索失败: {exc}", None)

    @app.post("/api/v1/entities/aliases")
    async def create_entity_alias(req: EntityAliasRequest):
        """Persist a user-confirmed alias for future entity resolution."""
        try:
            from dra_ma.tools.entity_resolver import EntityResolver

            resolver = EntityResolver()
            saved = await resolver.save_alias(
                alias=req.alias,
                canonical_name=req.canonical_name,
                kg_node_id=req.kg_node_id,
                entity_type=req.entity_type,
                source=req.source or "user_confirmed",
            )
            return _api_ok(saved, "alias_saved")
        except ValueError as exc:
            return _api_error(40002, str(exc), None)
        except Exception as exc:
            _logger.exception("[EntityAlias] failed: %s", exc)
            return _api_error(50004, f"实体别名保存失败: {exc}", None)

    # ── DRA-MA KG Q&A routes ──────────────────────────────────────

    @app.post("/api/v1/chat/recommend", response_model=ApiSuccessResponse)
    async def recommend(payload: ChatRecommendRequest, request: Request) -> ApiSuccessResponse:
        trace = TraceContext(
            sessionId=payload.session_id,
            roundId=payload.round_id,
            traceId=getattr(request.state, "trace_id", get_trace_id()),
        )
        result = await kg_system.handle_request(query=payload.query, history=payload.history, trace=trace)
        return ApiSuccessResponse(**result)

    @app.get("/api/v1/graph/expand")
    def expand_graph(id: str, type: str):
        return kg_system.expand_node(node_id=id, node_type=type)

    @app.get("/api/v1/chat/recommend-stream")
    async def recommend_stream(request: Request, query: str = "", history: str = "{}"):
        _logger.warning("[DEPRECATED] GET /recommend-stream called — forwarding to unified-stream intent_hint=graph_qa")

        session_id = request.query_params.get("sessionId", f"sess-{uuid4().hex[:10]}")
        try:
            round_id = int(request.query_params.get("roundId", "1"))
        except ValueError:
            round_id = 1

        req = UnifiedStreamRequest(
            query=query,
            sessionId=session_id,
            roundId=round_id,
            intentHint="graph_qa",
        )
        return await unified_stream(req, request)

    @app.post("/api/v1/chat/route", response_model=RouteResponse)
    async def route_intent(req: RouteRequest) -> RouteResponse:
        query = req.query
        risk_keywords = ["风险", "异常", "传导", "暴雷", "合规", "违规", "监管", "处罚", "事故", "损失"]
        graph_keywords = ["图谱", "查询", "关系", "关联", "路径", "公司", "企业", "人物", "事件", "节点"]
        risk_report_keywords = [
            "风险报告", "风险分析", "社区风险", "风险社区", "群体风险",
            "传导分析", "合规报告", "社区报告", "治理报告", "协同治理",
            "社区风险报告", "群体风险报告",
        ]

        has_risk = any(kw in query for kw in risk_keywords)
        has_graph = any(kw in query for kw in graph_keywords)
        has_risk_report = any(kw in query for kw in risk_report_keywords)

        # Risk report queries: explicit risk analysis / community risk intent
        if has_risk_report or (has_risk and not has_graph):
            return RouteResponse(route="risk", clarify_message=None)
        if has_graph or (has_risk and has_graph):
            return RouteResponse(route="graph", clarify_message=None)
        return RouteResponse(
            route="clarify",
            clarify_message="请问您是想查询图谱中的实体关系，还是进行协同治理分析？",
        )

    # ── Unified Stream (new main entry point) ───────────────────────

    class UnifiedStreamRequest(BaseModel):
        query: str = Field(default="", description="用户查询")
        file_content: str | None = Field(default=None, alias="fileContent")
        session_id: str = Field(default="", alias="sessionId")
        round_id: int = Field(default=1, alias="roundId")
        max_hop: int = Field(default=3, ge=1, le=5, alias="maxHop")
        intent_hint: str | None = Field(default=None, alias="intentHint")
        confirmed_entities: list[dict[str, Any]] = Field(default_factory=list, alias="confirmedEntities")
        workflow: str | None = Field(default=None)

    @app.post("/api/v1/chat/unified-stream")
    async def unified_stream(req: UnifiedStreamRequest, request: Request):
        """Unified SSE endpoint — single authoritative entry for both graph_qa and risk_analysis.

        Internal pipeline:
          IntentAgent → EntityResolution → DRAEngine → GraphAnalytics → RiskPlugins
        """
        _logger_sse = logging.getLogger("api.router.unified")

        async def event_generator():
            try:
                from dra_ma.orchestrator.unified_engine import UnifiedEngine
                unified = UnifiedEngine(dra_engine=kg_system, demo=False)

                async for line in unified.stream(
                    query=req.query,
                    session_id=req.session_id,
                    round_id=req.round_id,
                    max_hop=req.max_hop,
                    intent_hint=req.intent_hint,
                    file_content=req.file_content,
                    confirmed_entities=req.confirmed_entities,
                    workflow=req.workflow,
                ):
                    # UnifiedEngine yields JSON envelope strings — wrap as SSE
                    data = _json.loads(line) if isinstance(line, str) else line
                    event_type = data.get("type", "stage")
                    yield f"event: {event_type}\ndata: {_json.dumps(data, ensure_ascii=False)}\n\n"

            except Exception as exc:
                _logger_sse.exception("[UnifiedStream] Fatal error: %s", exc)
                yield f"event: error\ndata: {_json.dumps({'error': str(exc)}, ensure_ascii=False)}\n\n"
                yield f"event: done\ndata: {{}}\n\n"

        return StreamingResponse(
            event_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Access-Control-Allow-Origin": "*",
            },
        )

    @app.post("/api/v1/risk-paths")
    async def create_risk_paths(req: RiskPathsRequest, request: Request):
        """REST resource for risk transmission path analysis.

        This endpoint is a non-streaming facade over the governance pipeline.
        It returns candidate paths, interpreted paths, node paths, community
        context, anomaly findings, and graph highlight IDs.
        """
        try:
            query = _build_risk_path_query(req)
            collected = await _collect_unified_analysis(
                kg_system,
                query=query,
                session_id=req.session_id or getattr(request.state, "trace_id", ""),
                round_id=req.round_id,
                max_hop=req.max_hop,
                file_content=req.file_content,
            )
            risk_paths_payload = collected.get("risk_paths") or {}
            if isinstance(risk_paths_payload, list):
                interpreted_paths = risk_paths_payload
                candidate_paths = collected.get("candidate_risk_paths") or []
                merged_paths = interpreted_paths
            else:
                candidate_paths = risk_paths_payload.get("candidate_paths") or collected.get("candidate_risk_paths") or []
                interpreted_paths = risk_paths_payload.get("interpreted_paths") or []
                merged_paths = risk_paths_payload.get("merged_paths") or interpreted_paths or candidate_paths

            graph_highlight = {
                "node_ids": sorted({
                    str(node_id)
                    for path in merged_paths
                    for node_id in (path.get("node_ids") or path.get("nodeIds") or [])
                }),
                "edge_ids": sorted({
                    str(edge_id)
                    for path in merged_paths
                    for edge_id in (path.get("edge_ids") or path.get("edgeIds") or [])
                }),
            }
            return _api_ok({
                "query": query,
                "mode": req.mode,
                "candidate_paths": candidate_paths,
                "risk_paths": interpreted_paths,
                "merged_paths": merged_paths,
                "node_paths": merged_paths,
                "community_paths": [],
                "anomaly_findings": collected.get("anomaly_findings") or [],
                "community_info": collected.get("community"),
                "entity_community_map": collected.get("entity_community_map"),
                "subgraph": collected.get("subgraph"),
                "graph_highlight": graph_highlight,
            })
        except Exception as exc:
            _logger.exception("[RiskPaths] failed: %s", exc)
            return _api_error(50001, f"风险传导分析失败: {exc}", None)

    @app.post("/api/v1/governance/reports")
    async def create_governance_report(req: GovernanceReportRequest, request: Request):
        """REST resource for collaborative governance community reports."""
        try:
            collected = await _collect_unified_analysis(
                kg_system,
                query=req.query,
                session_id=req.session_id or getattr(request.state, "trace_id", ""),
                round_id=req.round_id,
                max_hop=req.max_hop,
                file_content=req.file_content,
            )
            report = collected.get("report")
            if not report:
                report = {
                    "report_id": f"WIND-RPT-{uuid4().hex[:8].upper()}",
                    "query_summary": req.query,
                    "executive_summary": "本轮分析未形成完整报告，已返回可确认的中间结果。",
                    "community_info": collected.get("community"),
                    "entity_community_map": collected.get("entity_community_map"),
                    "risk_paths": (collected.get("risk_paths") or {}).get("interpreted_paths", [])
                    if isinstance(collected.get("risk_paths"), dict) else [],
                    "anomaly_findings": collected.get("anomaly_findings") or [],
                    "compliance_matches": collected.get("compliance_matches") or [],
                    "risk_scores": collected.get("risk_scores"),
                    "governance_plan": collected.get("governance_plan"),
                    "recommendations": [],
                    "markdown_report": "",
                    "subgraph_summary": {
                        "node_count": len((collected.get("subgraph") or {}).get("nodes", [])),
                        "edge_count": len((collected.get("subgraph") or {}).get("edges", [])),
                    },
                }
            if not report.get("report_id"):
                report["report_id"] = f"WIND-RPT-{uuid4().hex[:8].upper()}"
            if not report.get("query_summary"):
                report["query_summary"] = req.query
            export_files = {
                fmt: None
                for fmt in req.export_formats
                if fmt in {"markdown", "docx", "pdf"}
            }
            if "markdown" in export_files:
                export_files["markdown"] = {
                    "filename": f"{report.get('report_id', 'governance-report')}.md",
                    "content": report.get("markdown_report") or report.get("integrated_report") or "",
                }
            return _api_ok({
                **report,
                "export_files": export_files,
                "source_events": {
                    "has_community": collected.get("community") is not None,
                    "has_risk_paths": collected.get("risk_paths") is not None,
                    "has_governance": collected.get("governance_plan") is not None,
                },
            })
        except Exception as exc:
            _logger.exception("[GovernanceReports] failed: %s", exc)
            return _api_error(50002, f"协同治理社区报告生成失败: {exc}", None)

    @app.post("/api/v1/chat/upload")
    async def upload_file(file: UploadFile = File(...)):
        MAX_CHARS = 50_000
        try:
            filename = file.filename or "unknown.txt"
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

            if ext in ("txt", "md", ""):
                raw = await file.read()
                # Try UTF-8 first, then common Chinese encodings
                for encoding in ("utf-8", "gbk", "gb2312", "gb18030"):
                    try:
                        text = raw.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    text = raw.decode("utf-8", errors="replace")

            elif ext == "pdf":
                import fitz  # PyMuPDF
                raw = await file.read()
                doc = fitz.open(stream=raw, filetype="pdf")
                parts = []
                for page in doc:
                    parts.append(page.get_text())
                doc.close()
                text = "\n".join(parts)

            elif ext == "docx":
                import io, docx
                raw = await file.read()
                doc = docx.Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs)

            else:
                return {
                    "success": False,
                    "message": f"不支持的文件类型: .{ext} (支持 .txt .md .docx .pdf)",
                }

            char_count = len(text)
            truncated = char_count > MAX_CHARS
            if truncated:
                text = text[:MAX_CHARS]

            return {
                "success": True,
                "data": {
                    "filename": filename,
                    "text": text,
                    "char_count": char_count,
                    "truncated": truncated,
                },
            }

        except Exception as exc:
            _logger.exception("File upload failed: %s", exc)
            return {"success": False, "message": f"文件解析失败: {str(exc)}"}

    @app.post("/api/v1/chat/analyze", response_class=StreamingResponse)
    async def analyze_risk(req: AnalyzeRequest):
        async def generate():
            try:
                engine = risk_engine
                if engine is None:
                    yield f"event: stage\ndata: {_json.dumps({'content': '协同治理引擎未初始化'}, ensure_ascii=False)}\n\n"
                    yield f"event: done\ndata: {{}}\n\n"
                    return

                async for update in engine.analyze_stream(query=req.query):
                    if "stage" in update:
                        stage_name = update["stage"]
                        content = update.get("content", "")
                        yield f"event: stage\ndata: {_json.dumps({'stage': stage_name, 'content': content}, ensure_ascii=False)}\n\n"
                    elif "entity_stats" in update:
                        yield f"event: entity_stats\ndata: {_json.dumps(update['entity_stats'], ensure_ascii=False)}\n\n"
                    elif "community" in update:
                        yield f"event: community\ndata: {_json.dumps(update['community'], ensure_ascii=False)}\n\n"
                    elif "risk_paths" in update:
                        yield f"event: risk_paths\ndata: {_json.dumps(update['risk_paths'], ensure_ascii=False)}\n\n"
                    elif "output" in update:
                        output = update['output']
                        yield f"event: analysis_text\ndata: {_json.dumps({'chunk': output.get('markdown_report', '')}, ensure_ascii=False)}\n\n"
                        ec = output.get('echarts_config')
                        rd = output.get('raw_data', [])
                        yield f"event: echarts_config\ndata: {_json.dumps(ec, ensure_ascii=False) if ec else 'null'}\n\n"
                        yield f"event: raw_data\ndata: {_json.dumps(rd, ensure_ascii=False)}\n\n"
                        yield f"event: done\ndata: {_json.dumps({'row_count': output.get('subtasks_completed', 0)}, ensure_ascii=False)}\n\n"

            except Exception as exc:
                _logger.exception("[AnalyzeSSE] Stream error: %s", exc)
                yield f"event: error\ndata: {_json.dumps({'error': str(exc)}, ensure_ascii=False)}\n\n"
                yield f"event: done\ndata: {{}}\n\n"

        return StreamingResponse(
            generate(),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Access-Control-Allow-Origin": "*"}
        )

    # ── Community-aware Risk Analysis SSE (for KnowledgeQA chat) ──

    async def _match_community_to_query(db, query: str, layer: str | None = None) -> dict | None:
        """Detect communities and return the one most relevant to the query."""
        try:
            from kg_query.analytics.graph_analytics import GraphAnalytics
            analytics = GraphAnalytics(db)
            result = analytics.detect_communities(
                layer=layer, method="louvain", min_community_size=3
            )
            communities = result.get("communities", [])
            if not communities:
                return None

            # Extract keywords from query (simple character-level segmentation)
            query_chars = set(query.replace(" ", ""))
            # Score each community by how many top-entity names overlap with query
            best = None
            best_score = 0
            for comm in communities:
                score = 0
                for ent in comm.get("top_entities", []):
                    name = ent.get("name", "")
                    # Count overlapping characters
                    name_chars = set(name)
                    overlap = len(query_chars & name_chars)
                    score += overlap
                if score > best_score:
                    best_score = score
                    best = comm

            if best and best_score > 0:
                return {
                    "communities": [best],
                    "algorithm": "louvain_match",
                    "matched_community_id": best["community_id"],
                }
            return None
        except Exception as exc:
            _logger.warning("[RiskStream] Community matching failed: %s", exc)
            return None

    async def _build_risk_stream_events(
        engine,
        query: str,
        focus_entities: list[str],
        parsed_max_hop: int,
        parsed_community_id: int | None,
        file_content: str | None,
        _logger_sse,
    ):
        """Shared risk-stream event generator used by both GET and POST endpoints."""
        if engine is None:
            yield f"event: stage\ndata: {_json.dumps({'stage': 'planning', 'content': '协同治理引擎未初始化'}, ensure_ascii=False)}\n\n"
            yield f"event: done\ndata: {{}}\n\n"
            return

        try:
            effective_query = query
            effective_focus_entities = list(focus_entities)

            # If file content provided and no focus_entities, extract them
            if file_content and not effective_focus_entities:
                yield f"event: stage\ndata: {_json.dumps({'stage': 'planning', 'content': '正在分析文件内容，提取关键实体...'}, ensure_ascii=False)}\n\n"
                extraction = await engine._extract_from_file_content_llm(file_content)

                if extraction.get("is_financial_risk_relevant", True):
                    effective_focus_entities = extraction.get("entities", [])
                    risk_signals = extraction.get("risk_signals", [])
                    summary = extraction.get("summary", "")

                    if summary and not query.strip():
                        effective_query = f"请分析以下文档内容的协同治理情况: {summary}"
                    if risk_signals:
                        effective_query += f"\n\n文档中识别到的风险信号: {'; '.join(risk_signals)}"

                    yield f"event: stage\ndata: {_json.dumps({'stage': 'planning', 'content': f'提取到 {len(effective_focus_entities)} 个关键实体: {", ".join(effective_focus_entities[:8])}'}, ensure_ascii=False)}\n\n"
                else:
                    yield f"event: stage\ndata: {_json.dumps({'stage': 'planning', 'content': '文档未识别到明显的金融风控相关内容，将进行通用分析'}, ensure_ascii=False)}\n\n"

            # Phase A: Community detection & selection
            matched_community: dict | None = None

            if parsed_community_id is not None and hasattr(engine, "_db"):
                yield f"event: stage\ndata: {_json.dumps({'stage': 'retrieving', 'content': f'获取社区 #{parsed_community_id} 子图...'}, ensure_ascii=False)}\n\n"
                try:
                    from kg_query.analytics.graph_analytics import GraphAnalytics
                    analytics = GraphAnalytics(engine._db)
                    sub = analytics.get_community_subgraph(parsed_community_id, limit=200)
                    nodes = sub.get("nodes", [])
                    community_focus = [
                        n.get("name") or n.get("properties", {}).get("name", "")
                        for n in nodes
                        if (n.get("name") or n.get("properties", {}).get("name", ""))
                    ][:20]
                    if community_focus:
                        effective_focus_entities = community_focus
                    if nodes:
                        yield f"event: subgraph\ndata: {_json.dumps({'nodes': nodes, 'edges': sub.get('edges', [])}, ensure_ascii=False)}\n\n"
                    yield f"event: community\ndata: {_json.dumps({'community_id': parsed_community_id, 'size': len(nodes), 'top_entities': [{'id': n.get('id',''), 'name': n.get('name','') or n.get('properties',{}).get('name',''), 'label': (n.get('labels') or [''])[0]} for n in nodes[:5]]}, ensure_ascii=False)}\n\n"
                except Exception as exc:
                    _logger_sse.warning("[RiskStream] Community subgraph fetch failed: %s", exc)

            elif hasattr(engine, "_db"):
                yield f"event: stage\ndata: {_json.dumps({'stage': 'retrieving', 'content': '检测图谱社区结构...'}, ensure_ascii=False)}\n\n"
                matched_community = await _match_community_to_query(engine._db, effective_query)
                if matched_community:
                    yield f"event: community\ndata: {_json.dumps(matched_community, ensure_ascii=False)}\n\n"
                    matched_entities = [
                        ent.get("name", "") for ent in matched_community.get("communities", [[]])[0].get("top_entities", [])
                    ]
                    if matched_entities:
                        effective_focus_entities = matched_entities
                    matched_cid = matched_community.get("matched_community_id", matched_community.get("communities", [{}])[0].get("community_id", "?"))
                    matched_size = matched_community.get("communities", [{}])[0].get("size", 0)
                    yield f"event: stage\ndata: {_json.dumps({'stage': 'retrieving', 'content': f'匹配到社区 #{matched_cid} ({matched_size} 个节点)，开始协同治理分析...'}, ensure_ascii=False)}\n\n"

            # Phase B: Run the 5-agent risk analysis pipeline
            async for update in engine.analyze_stream(
                query=effective_query,
                focus_entities=effective_focus_entities,
                max_hop=parsed_max_hop,
            ):
                if "stage" in update:
                    stage_name = update["stage"]
                    if stage_name == "subgraph":
                        sub_data = {
                            "nodes": update.get("nodes", []),
                            "edges": update.get("edges", []),
                        }
                        yield f"event: subgraph\ndata: {_json.dumps(sub_data, ensure_ascii=False)}\n\n"
                    else:
                        yield f"event: stage\ndata: {_json.dumps({'stage': stage_name, 'content': update.get('content', '')}, ensure_ascii=False)}\n\n"
                elif "output" in update:
                    output = update["output"]
                    _save_report(output, effective_query)
                    yield f"event: report\ndata: {_json.dumps(output, ensure_ascii=False)}\n\n"
                    yield f"event: done\ndata: {{}}\n\n"

        except Exception as exc:
            _logger_sse.exception("[RiskStream] Stream error: %s", exc)
            yield f"event: error\ndata: {_json.dumps({'error': str(exc)}, ensure_ascii=False)}\n\n"
            yield f"event: done\ndata: {{}}\n\n"

    @app.get("/api/v1/chat/risk-stream")
    async def chat_risk_stream(
        request: Request,
        query: str = "",
        sessionId: str = "",
        roundId: str = "1",
        communityId: str = "",
        maxHop: str = "3",
    ):
        _logger.warning("[DEPRECATED] GET /risk-stream called — forwarding to unified-stream intent_hint=risk_analysis")

        try:
            parsed_max_hop = int(maxHop)
        except ValueError:
            parsed_max_hop = 3

        try:
            round_id = int(roundId)
        except ValueError:
            round_id = 1

        req = UnifiedStreamRequest(
            query=query,
            sessionId=sessionId,
            roundId=round_id,
            maxHop=parsed_max_hop,
            intentHint="risk_analysis",
        )
        return await unified_stream(req, request)

    @app.post("/api/v1/chat/risk-stream")
    async def chat_risk_stream_post(req_in: RiskStreamRequest, request: Request):
        _logger.warning("[DEPRECATED] POST /risk-stream called — forwarding to unified-stream intent_hint=risk_analysis")

        req = UnifiedStreamRequest(
            query=req_in.query,
            fileContent=req_in.file_content,
            sessionId=req_in.session_id,
            roundId=req_in.round_id,
            maxHop=req_in.max_hop,
            intentHint="risk_analysis",
        )
        return await unified_stream(req, request)

    # ── Chat history persistence ───────────────────────────────────

    class ChatHistoryRecord(BaseModel):
        session_id: str = Field(..., alias="sessionId")
        role: str = Field(default="user")
        content: str = Field(default="")
        message_id: str = Field(default_factory=lambda: f"msg-{uuid4().hex[:10]}", alias="messageId")

    @app.get("/api/v1/chat/history/{session_id}")
    async def get_chat_history(session_id: str):
        """Retrieve chat history for a session from Neo4j."""
        if risk_engine is None or not hasattr(risk_engine, "_db"):
            return {"success": True, "data": {"messages": [], "session_id": session_id}}
        try:
            db = risk_engine._db
            cypher = """
            MATCH (s:ChatSession {session_id: $sid})-[:CONTAINS]->(m:ChatMessage)
            RETURN m ORDER BY m.timestamp ASC
            """
            rows, _ = db.execute_read_with_summary(cypher, {"sid": session_id})
            messages = []
            for row in rows:
                msg = row.get("m", {})
                props = msg.get("properties", msg) if isinstance(msg, dict) else {}
                messages.append({
                    "id": props.get("message_id", ""),
                    "role": props.get("role", "user"),
                    "content": props.get("content", ""),
                    "timestamp": str(props.get("timestamp", "")),
                })
            return {"success": True, "data": {"messages": messages, "session_id": session_id}}
        except Exception as exc:
            _logger.warning("[ChatHistory] GET failed: %s", exc)
            return {"success": True, "data": {"messages": [], "session_id": session_id}}

    @app.post("/api/v1/chat/history")
    async def save_chat_history(records: list[ChatHistoryRecord]):
        """Store chat messages in Neo4j for server-side session persistence."""
        if risk_engine is None or not hasattr(risk_engine, "_db"):
            return {"success": False, "message": "引擎未初始化"}
        try:
            db = risk_engine._db
            for rec in records:
                cypher = """
                MERGE (s:ChatSession {session_id: $sid})
                CREATE (m:ChatMessage {
                    message_id: $mid,
                    role: $role,
                    content: $content,
                    timestamp: datetime()
                })
                CREATE (s)-[:CONTAINS]->(m)
                """
                db.execute_read(cypher, {
                    "sid": rec.session_id,
                    "mid": rec.message_id,
                    "role": rec.role,
                    "content": rec.content[:2000],
                })
            return {"success": True, "message": f"已保存 {len(records)} 条消息"}
        except Exception as exc:
            _logger.exception("[ChatHistory] POST failed: %s", exc)
            return {"success": False, "message": str(exc)}

    # ── Risk Analysis routes ──────────────────────────────────────

    @app.get("/api/v1/risk/analyze-stream")
    async def risk_analyze_stream(request: Request, query: str = "", focus_entities: str = "[]"):
        _logger_sse = logging.getLogger("api.router.risk")

        try:
            parsed_entities: list[str] = _json.loads(focus_entities) or []
        except Exception:
            parsed_entities = []

        async def event_generator():
            try:
                engine = risk_engine
                if engine is None:
                    yield f"event: stage\ndata: {_json.dumps({'content': '协同治理引擎未初始化'}, ensure_ascii=False)}\n\n"
                    yield f"event: done\ndata: {{}}\n\n"
                    return

                async for update in engine.analyze_stream(
                    query=query,
                    focus_entities=parsed_entities,
                ):
                    if "stage" in update:
                        stage_name = update["stage"]
                        if stage_name == "subgraph":
                            sub_data = {
                                "nodes": update.get("nodes", []),
                                "edges": update.get("edges", []),
                            }
                            yield f"event: subgraph\ndata: {_json.dumps(sub_data, ensure_ascii=False)}\n\n"
                        else:
                            yield f"event: stage\ndata: {_json.dumps({'stage': stage_name, 'content': update.get('content', '')}, ensure_ascii=False)}\n\n"
                    elif "entity_stats" in update:
                        yield f"event: entity_stats\ndata: {_json.dumps(update['entity_stats'], ensure_ascii=False)}\n\n"
                    elif "community" in update:
                        yield f"event: community\ndata: {_json.dumps(update['community'], ensure_ascii=False)}\n\n"
                    elif "risk_paths" in update:
                        yield f"event: risk_paths\ndata: {_json.dumps(update['risk_paths'], ensure_ascii=False)}\n\n"
                    elif "output" in update:
                        output = update["output"]
                        report_id = _save_report(output, query) or f"RPT-{uuid4().hex[:8].upper()}"
                        output.setdefault("report_id", report_id)
                        output.setdefault("generated_at", __import__("datetime").datetime.now().isoformat())
                        output.setdefault("query_summary", query[:200] if query else "")
                        output.setdefault("legal_basis", [])
                        output.setdefault("penalty_cases", [])
                        yield f"event: report\ndata: {_json.dumps(output, ensure_ascii=False)}\n\n"
                        yield f"event: done\ndata: {{}}\n\n"

            except Exception as exc:
                _logger_sse.exception("[RiskSSE] Stream error: %s", exc)
                yield f"event: error\ndata: {_json.dumps({'error': str(exc)}, ensure_ascii=False)}\n\n"
                yield f"event: done\ndata: {{}}\n\n"

        return StreamingResponse(
            event_generator(),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Access-Control-Allow-Origin": "*"},
        )

    @app.post("/api/v1/risk/analyze")
    async def risk_analyze(req: RiskAnalyzeRequest):
        if risk_engine is None:
            return {"success": False, "message": "协同治理引擎未初始化"}
        result = await risk_engine.analyze(
            query=req.query,
            focus_entities=req.focus_entities,
            max_hop=req.max_hop,
            trigger_event=req.trigger_event,
        )
        return {"success": True, "data": result}

    # ── Report persistence (Neo4j RiskReport nodes) ─────────────────

    def _save_report(output: dict, query: str) -> str | None:
        """Persist analysis result as a RiskReport node in Neo4j. Returns report_id."""
        if risk_engine is None or not hasattr(risk_engine, "_db"):
            return None
        try:
            report_id = f"RPT-{uuid4().hex[:8].upper()}"
            db = risk_engine._db
            cypher = """
            CREATE (r:RiskReport {
                report_id: $report_id,
                query: $query,
                executive_summary: $summary,
                overall_risk_level: $risk_level,
                risk_path_count: $path_count,
                anomaly_count: $anomaly_count,
                compliance_count: $compliance_count,
                node_count: $node_count,
                edge_count: $edge_count,
                created_at: datetime()
            })
            """
            sub = output.get("subgraph_summary", {}) or {}
            db.execute_read(
                cypher,
                {
                    "report_id": report_id,
                    "query": query[:500],
                    "summary": output.get("executive_summary", "")[:500],
                    "risk_level": output.get("overall_risk_level", "medium"),
                    "path_count": len(output.get("risk_paths", [])),
                    "anomaly_count": len(output.get("anomaly_findings", [])),
                    "compliance_count": len(output.get("compliance_matches", [])),
                    "node_count": sub.get("node_count", 0),
                    "edge_count": sub.get("edge_count", 0),
                },
            )
            logger_save = logging.getLogger("api.router.report")
            logger_save.info("Saved report %s to Neo4j", report_id)
            return report_id
        except Exception as exc:
            logger_save = logging.getLogger("api.router.report")
            logger_save.warning("Failed to save report: %s", exc)
            return None

    @app.get("/api/v1/risk/reports")
    async def list_reports(page: int = 1, risk_level: str = "", limit: int = 20):
        """List saved risk reports from Neo4j."""
        if risk_engine is None:
            return {"success": True, "data": {"reports": [], "total": 0, "page": page}}
        try:
            db = risk_engine._db
            level_filter = "WHERE r.overall_risk_level = $risk_level" if risk_level else ""
            skip = (page - 1) * limit
            cypher = f"""
            MATCH (r:RiskReport) {level_filter}
            RETURN r ORDER BY r.created_at DESC SKIP $skip LIMIT $limit
            """
            count_cypher = "MATCH (r:RiskReport) RETURN count(r) AS total"
            rows, _ = db.execute_read_with_summary(cypher, {"risk_level": risk_level, "skip": skip, "limit": limit})
            count_rows, _ = db.execute_read_with_summary(count_cypher)
            total = count_rows[0].get("total", 0) if count_rows else 0

            reports = []
            for row in rows:
                data = row.get("r", {})
                props = data.get("properties", data) if isinstance(data, dict) else {}
                reports.append({
                    "report_id": props.get("report_id", ""),
                    "query": props.get("query", ""),
                    "executive_summary": props.get("executive_summary", "")[:200],
                    "overall_risk_level": props.get("overall_risk_level", "medium"),
                    "risk_path_count": props.get("risk_path_count", 0),
                    "anomaly_count": props.get("anomaly_count", 0),
                    "compliance_count": props.get("compliance_count", 0),
                    "created_at": str(props.get("created_at", "")),
                })
            return {"success": True, "data": {"reports": reports, "total": total, "page": page}}
        except Exception as exc:
            return {"success": False, "message": f"Failed to list reports: {exc}", "data": {"reports": [], "total": 0, "page": page}}

    @app.post("/api/v1/risk/reports/export-docx")
    async def export_report_docx(req: ReportDocxExportRequest):
        """Export the current risk report as a real DOCX document."""
        try:
            from dra_ma.reporting import DocxExporter

            exporter = DocxExporter()
            report = req.report or {}
            report_id = req.report_id or report.get("report_id") or f"WIND-RPT-{uuid4().hex[:8]}"
            docx_bytes = exporter.export(
                report,
                {
                    "report_id": report_id,
                    "query_text": req.query_text or report.get("query_summary") or "",
                    "generated_at": report.get("generated_at") or "",
                },
            )
            safe_name = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in report_id)
            return Response(
                content=docx_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename={safe_name}.docx",
                    "Cache-Control": "no-store",
                },
            )
        except Exception as exc:
            _logger.exception("[ReportExport] DOCX export failed: %s", exc)
            return {"success": False, "message": f"DOCX 导出失败: {exc}"}

    @app.get("/api/v1/risk/reports/{report_id}")
    async def get_report(report_id: str):
        """Get a single saved report by ID."""
        if risk_engine is None:
            return {"success": False, "message": "引擎未初始化"}
        try:
            db = risk_engine._db
            cypher = "MATCH (r:RiskReport {report_id: $id}) RETURN r"
            rows, _ = db.execute_read_with_summary(cypher, {"id": report_id})
            if not rows:
                return {"success": False, "message": "Report not found"}
            props = rows[0].get("r", {})
            props = props.get("properties", props) if isinstance(props, dict) else {}
            return {"success": True, "data": {
                "report_id": props.get("report_id", ""),
                "query": props.get("query", ""),
                "executive_summary": props.get("executive_summary", ""),
                "overall_risk_level": props.get("overall_risk_level", "medium"),
                "risk_path_count": props.get("risk_path_count", 0),
                "anomaly_count": props.get("anomaly_count", 0),
                "compliance_count": props.get("compliance_count", 0),
                "node_count": props.get("node_count", 0),
                "edge_count": props.get("edge_count", 0),
                "created_at": str(props.get("created_at", "")),
            }}
        except Exception as exc:
            return {"success": False, "message": str(exc)}

    # ── File Upload ───────────────────────────────────────────

    @app.post("/api/v1/chat/upload")
    async def upload_file(request: Request):
        """Upload a text file (.txt/.md/.docx/.pdf) and return extracted text."""
        import os
        import tempfile

        ALLOWED_EXTENSIONS = {'.txt', '.md', '.docx', '.pdf'}
        BLOCKED_EXTENSIONS = {'.exe', '.bat', '.sh', '.com', '.dll', '.js', '.vbs', '.ps1', '.scr', '.msi', '.pif', '.cmd', '.cpl'}
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        MAX_TEXT_LENGTH = 50000

        content_type = request.headers.get("content-type", "")
        if "multipart/form-data" not in content_type:
            return {"success": False, "message": "请求格式错误，需要 multipart/form-data"}

        form = await request.form()
        file = form.get("file")
        if file is None:
            return {"success": False, "message": "未找到上传文件"}

        filename = getattr(file, 'filename', 'unknown')
        ext = os.path.splitext(filename)[1].lower()

        if ext in BLOCKED_EXTENSIONS:
            return {"success": False, "message": f"不支持的文件类型: {ext}"}

        if ext not in ALLOWED_EXTENSIONS:
            return {"success": False, "message": f"不支持的文件格式: {ext}，仅支持 .txt .md .docx .pdf"}

        content = await file.read()

        if len(content) > MAX_FILE_SIZE:
            return {"success": False, "message": "文件过大（最大 10MB），请拆分后重试"}

        # Magic number validation
        if ext == '.pdf' and content[:4] != b'%PDF':
            return {"success": False, "message": "文件内容与扩展名不匹配（非有效 PDF 文件）"}
        if ext == '.docx' and content[:4] != b'PK\x03\x04':
            return {"success": False, "message": "文件内容与扩展名不匹配（非有效 DOCX 文件）"}

        text = ""
        try:
            if ext in ('.txt', '.md'):
                try:
                    text = content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        text = content.decode('gbk')
                    except Exception:
                        return {"success": False, "message": "文件编码不支持，请使用 UTF-8 或 GBK 编码"}
            elif ext == '.pdf':
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                try:
                    from data_collection.file_import.pdf_parser import parse_pdf_hybrid
                    text = parse_pdf_hybrid(tmp_path) or ""
                finally:
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass
            elif ext == '.docx':
                try:
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(content))
                    text = "\n".join(p.text for p in doc.paragraphs if p.text)
                except ImportError:
                    return {"success": False, "message": "python-docx 未安装，无法解析 .docx 文件"}
                except Exception as e:
                    return {"success": False, "message": f"DOCX 解析失败: {str(e)}"}
        except Exception as e:
            return {"success": False, "message": f"文件解析失败: {str(e)}"}

        if not text or not text.strip():
            return {"success": False, "message": "文件内容为空或无法提取文本"}

        truncated = len(text) > MAX_TEXT_LENGTH
        if truncated:
            text = text[:MAX_TEXT_LENGTH]

        return {
            "success": True,
            "data": {
                "filename": filename,
                "text": text,
                "char_count": len(text),
                "truncated": truncated,
            },
        }

    # ═══════════════════════════════════════════════════════════════
    # End Module C routes
    # ═══════════════════════════════════════════════════════════════
