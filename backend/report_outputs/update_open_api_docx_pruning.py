from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DOC_PATH = Path("docs/开放API接口文档_补全版.docx")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        for paragraph in hdr[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Microsoft YaHei"
                run.font.bold = True
                run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(9)


def main() -> None:
    doc = Document(DOC_PATH)

    add_heading(doc, "附录：图谱查询上限与智能剪枝策略", 1)
    add_para(
        doc,
        "本节适用于 POST /api/v1/graph/search-all 与 POST /api/v1/graph/expand/{node_id}。"
        "当 depth 大于 1 或中心节点连接大量银行、基金、证券、监管机构、时间、公共法规等高出度节点时，"
        "接口不应仅依赖 limit/nodeLimit 进行结果截断，而应在遍历阶段启用度数感知与证据优先剪枝，"
        "尽量返回对风险传导、群体发现和社区报告更有效的图谱。",
    )

    add_heading(doc, "A.1 剪枝原则", 2)
    for item in [
        "中心节点必须保留。",
        "一跳强关系节点优先保留，包括 GUARANTEE、CONTROLLER、CONTROL、INVEST、CAUSE、TRIGGERS。",
        "Event、Feature、Regulation 证据层优先保留，并优先形成可解释风险链路。",
        "一跳节点度数超过 maxExpandDegree 时，默认保留该节点及一跳关系，但不再将其加入下一跳 frontier。",
        "BANK、PFUND、PFCOMPANY、SECURITY、REGULATOR、EXCHANGE、TIME、Law、Regulation 等标签默认视为高出度低区分度候选。",
        "limit/nodeLimit 仍作为安全上限；若智能剪枝后仍达到上限，响应必须通过 warnings 与 summary.pruning 说明。",
    ]:
        add_para(doc, f"• {item}")

    add_heading(doc, "A.2 新增 Request 字段", 2)
    add_table(
        doc,
        ["字段", "类型", "默认值", "说明"],
        [
            ["traversalMode", "string", "bfs", "遍历模式：bfs/cascade。治理链路建议使用 cascade。"],
            ["prunePolicy", "string", "degree_aware", "剪枝策略：none/degree_aware/evidence_first。"],
            ["maxExpandDegree", "number", "200", "节点度数超过该值时默认不进入下一跳。"],
            ["highDegreeLabels", "string[]", "内置列表", "自定义高出度低区分度标签。为空时使用默认列表。"],
            ["keepHighDegreeNode", "boolean", "true", "是否保留高出度节点本身及一跳关系。"],
            ["includePruningSummary", "boolean", "true", "是否在 summary.pruning 中返回剪枝摘要。"],
        ],
    )

    add_heading(doc, "A.3 推荐请求示例", 2)
    add_para(
        doc,
        '{ "depth": 2, "limit": 500, "traversalMode": "cascade", '
        '"prunePolicy": "degree_aware", "maxExpandDegree": 200, '
        '"keepHighDegreeNode": true, "includePruningSummary": true }',
    )

    add_heading(doc, "A.4 Response 剪枝摘要", 2)
    add_para(
        doc,
        "两个图谱查询 API 的响应 summary 中建议包含 pruning 对象，用于说明剪枝策略、阈值、"
        "被保留为 terminal 的高出度节点数量、阻断的扩展次数以及按原因统计的剪枝数量。",
    )
    add_table(
        doc,
        ["字段", "类型", "说明"],
        [
            ["summary.pruning.policy", "string", "实际使用的剪枝策略。"],
            ["summary.pruning.maxExpandDegree", "number", "本次请求使用的最大可扩展度数阈值。"],
            ["summary.pruning.terminalHubCount", "number", "被保留为 terminal 且不继续扩展的高出度节点数量。"],
            ["summary.pruning.blockedExpansionCount", "number", "被阻断的下一跳扩展次数。"],
            ["summary.pruning.blockedByReason", "object", "按 high_degree、low_signal_label、layer_budget 等原因统计。"],
            ["summary.pruning.terminalHubs", "array", "最多返回若干高出度 terminal 节点示例，含 id/name/degree/labels/reason。"],
        ],
    )

    add_heading(doc, "A.5 Warning Code", 2)
    add_table(
        doc,
        ["Warning Code", "含义"],
        [
            ["PRUNING_APPLIED", "已启用智能剪枝策略。"],
            ["HIGH_DEGREE_NODE_TERMINATED", "高出度节点已保留为一跳 terminal，但不继续进入下一跳。"],
            ["LOW_SIGNAL_HUB_FILTERED", "低区分度高出度节点被过滤或仅保留摘要。"],
            ["NODE_LIMIT_REACHED_AFTER_PRUNING", "智能剪枝后仍达到节点上限。"],
            ["EDGE_LIMIT_REACHED_AFTER_PRUNING", "智能剪枝后仍达到关系上限。"],
        ],
    )

    add_heading(doc, "A.6 与两个图谱 API 的关系", 2)
    add_para(
        doc,
        "POST /api/v1/graph/search-all 用于从关键词找到图谱入口；POST /api/v1/graph/expand/{node_id} "
        "用于在用户确认节点后继续展开。若两者均使用 traversalMode=cascade、prunePolicy=degree_aware、"
        "maxExpandDegree=200、相同 depth 和 limit，则返回的业务子图应保持一致或高度一致。",
    )

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
