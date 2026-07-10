from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


DOC_PATH = Path("docs/开放API接口文档_补全版.docx")
MARKER = "A.7 limit 默认值与按跳数放大规则"


def set_font(paragraph, size: float = 10.5, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(size)
        run.font.bold = bold


def main() -> None:
    doc = Document(DOC_PATH)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in all_text:
        return

    heading = doc.add_heading(MARKER, level=2)
    set_font(heading, 13, True)

    paragraphs = [
        "POST /api/v1/graph/search-all 与 POST /api/v1/graph/expand/{node_id} 的 limit 默认值均为 1000。",
        "limit 表示单跳基础节点上限；当调用方未显式传入 nodeLimit 时，接口按 depth 自动计算有效节点上限：effectiveNodeLimit = min(limit × depth, 5000)。",
        "默认情况下，depth=1 返回上限为 1000，depth=2 返回上限为 2000，depth=3 返回上限为 3000，依此类推；当前接口 depth 最大为 5，因此默认最高有效节点上限为 5000。",
        "若调用方需要严格控制总节点数，可传入 nodeLimit。nodeLimit 优先级高于 limit × depth，将作为本次请求的绝对节点上限。",
        "edgeLimit 未传入时，默认按 effectiveNodeLimit × 4 计算，且不低于 2000。",
    ]
    for text in paragraphs:
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)
        set_font(paragraph)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["depth", "limit 默认值", "有效节点上限", "说明"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    rows = [
        ["1", "1000", "1000", "一跳默认最多返回 1000 个节点"],
        ["2", "1000", "2000", "二跳默认最多返回 2000 个节点"],
        ["3", "1000", "3000", "三跳默认最多返回 3000 个节点"],
        ["N", "1000", "min(1000 × N, 5000)", "按跳数线性放大，并受接口最大深度约束"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_font(paragraph, 9)

    doc.save(DOC_PATH)


if __name__ == "__main__":
    main()
